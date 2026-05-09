# -*- coding: utf-8 -*-
import json
import logging
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ==========================================
# 视觉风格定义
# ==========================================
if HAS_DOCX:
    ACCENT_COLOR = RGBColor(0, 0, 0) # 标准学术黑
else:  # pragma: no cover
    ACCENT_COLOR = None
LIGHT_FILL_HEX = 'F2F2F2'
BORDER_COLOR_HEX = '000000'

logger = logging.getLogger("P_Layer_Word")

class WordWriter:
    """
    P-Layer: 专业展示层 (Word 文档生成)
    负责将学术生成层 (G-Layer) 和 索引层 (K-Layer) 的产物转换为排版精美的人工验收文档。
    """
    
    def __init__(self, output_path: str | Path):
        self.output_path = Path(output_path)
        self.doc: Optional[Document] = None
        self.bookmark_id = 0
        self.image_cache: Dict[str, str] = {}

    def setup(self):
        """初始化文档并设置基础样式"""
        if not HAS_DOCX:
            raise RuntimeError("Missing 'python-docx' library. Please install it to use WordWriter.")
        
        self.doc = Document()
        sec = self.doc.sections[0]
        # 设置标准页边距
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        
        # 默认字体与对齐
        styles = self.doc.styles
        normal_font = styles['Normal'].font
        normal_font.name = 'Times New Roman'
        styles['Normal']._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
        normal_font.size = Pt(10.5)
        styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        styles['Normal'].paragraph_format.line_spacing = 1.15

        # 自定义专业样式 (v40.3 期刊版)
        for name, base, size, bold, color in [
            ('Pack Title', 'Title', 18, True, RGBColor(0, 0, 0)),
            ('Pack Subtitle', 'Subtitle', 10, False, RGBColor(50, 50, 50)),
            ('Pack Heading 1', 'Heading 1', 12, True, RGBColor(0, 0, 0)),
            ('Pack Heading 2', 'Heading 2', 11, True, RGBColor(0, 0, 0)),
            ('Pack Note', 'Normal', 9, False, RGBColor(0, 0, 0)),
            ('Pack Caption', 'Normal', 8.5, False, RGBColor(0, 0, 0)),
        ]:
            if name not in styles:
                st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                st.base_style = styles[base]
                st.font.name = 'Times New Roman'
                st._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                st.font.size = Pt(size)
                st.font.bold = bold
                st.font.color.rgb = color

    def _set_cell_shading(self, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tc_pr.append(shd)

    def _set_cell_margins(self, cell, top=80, start=100, bottom=80, end=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in('w:tcMar')
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = tcMar.find(qn(f'w:{m}'))
            if node is None:
                node = OxmlElement(f'w:{m}')
                tcMar.append(node)
            node.set(qn('w:w'), str(v))
            node.set(qn('w:type'), 'dxa')

    def _add_journal_visuals(self, items: List[Dict[str, Any]], is_table: bool = False):
        """ 紧凑型期刊风图文集成 (v40.4: 单栏紧凑) """
        for item in items:
            img_item = item.get('page_crop_image')
            img_path = img_item.get('image_path') if img_item else None
            
            if img_path and Path(img_path).exists():
                try:
                    safe_path = self._normalize_image(Path(img_path))
                    # 单栏下恢复较宽显示
                    width = self._get_image_width(safe_path, max_inches=5.8)
                    self.doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.doc.paragraphs[-1].add_run().add_picture(str(safe_path), width=Inches(width))
                    
                    cap = self.doc.add_paragraph(style='Pack Caption')
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    prefix = "Table" if is_table else "Figure"
                    num = item.get('figure_number') or item.get('table_number', '?')
                    cap.add_run(f"{prefix} {num}: ")
                    cap.add_run(item.get('caption', ''))
                    # 紧凑排版：缩小下方间距
                    cap.paragraph_format.space_after = Pt(2)
                except Exception as e:
                    self.doc.add_paragraph(f"[Image Error: {e}]", style='Pack Note')

    def _add_bookmark(self, paragraph, name: str):
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), str(self.bookmark_id))
        start.set(qn('w:name'), name)
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), str(self.bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
        self.bookmark_id += 1

    def _normalize_image(self, image_path: Path) -> Path:
        """纠正图像色彩空间，防止 Word 无法显示"""
        if not HAS_PIL:
            return image_path
        key = str(image_path.resolve())
        if key in self.image_cache:
            return Path(self.image_cache[key])
            
        try:
            with Image.open(image_path) as img:
                if img.mode not in ("RGB", "RGBA"):
                    img = img.convert("RGB")
                    out_dir = Path(tempfile.gettempdir()) / "p_layer_word_cache"
                    out_dir.mkdir(parents=True, exist_ok=True)
                    out_path = out_dir / f"{image_path.stem}_fixed.png"
                    img.save(out_path, format="PNG")
                    self.image_cache[key] = str(out_path)
                    return out_path
        except Exception as e:
            logger.warning(f"Image normalization failed for {image_path}: {e}")
        return image_path

    def _get_image_width(self, image_path: Path, max_inches: float = 6.2) -> float:
        """根据长宽比估算最佳显示宽度"""
        if not HAS_PIL:
            return max_inches
        try:
            with Image.open(image_path) as img:
                w, h = img.size
                aspect = w / max(h, 1)
                if aspect >= 1.6: return max_inches
                if aspect >= 1.1: return min(max_inches, 5.8)
                return min(max_inches, 4.8)
        except:
            return max_inches

    def add_title_page(self, title: str, subtitle: str = "", metadata: Dict[str, Any] = None):
        """添加首页导读与元数据"""
        doc = self.doc
        p = doc.add_paragraph(style='Pack Title')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title)

        if subtitle:
            p = doc.add_paragraph(style='Pack Subtitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(subtitle)

        doc.add_paragraph('')
        if metadata:
            table = doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for k, v in metadata.items():
                row = table.add_row().cells
                row[0].width = Cm(4)
                row[1].width = Cm(10)
                row[0].text = str(k)
                row[1].text = str(v)
                self._set_cell_shading(row[0], LIGHT_FILL_HEX)
                for cell in row:
                    self._set_cell_margins(cell)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_writing_points(self, points: List[Dict[str, Any]], figure_map: Dict[str, Dict[str, Any]] = None):
        """添加写作点区块"""
        doc = self.doc
        doc.add_page_break()
        doc.add_paragraph('核心写作点', style='Pack Heading 1')
        
        for idx, wp in enumerate(points, start=1):
            title = f"写作点 {idx:02d} | {wp.get('point_type', '结论')}"
            p = doc.add_paragraph(style='Pack Heading 2')
            p.add_run(title)
            self._add_bookmark(p, f"wp_{wp.get('writing_point_id', idx)}")

            # 主张 (Claim) - 加粗突出
            claim_p = doc.add_paragraph()
            claim_p.add_run(wp.get('claim', wp.get('representative_claim', ''))).bold = True
            
            # 元数据表格
            meta = doc.add_table(rows=0, cols=2)
            meta.autofit = False
            rows = [
                ('证据边界', f"{wp.get('boundary_type', '')} - {wp.get('boundary_note', '')}"),
                ('关联证据', wp.get('evidence_summary', '')),
                ('所处页码', str(wp.get('pages', wp.get('page', '')))),
                ('原始引用', " ".join(wp.get('original_reference_markers', []))),
            ]
            for label, val in rows:
                if not val or val == ' - ': continue
                cells = meta.add_row().cells
                cells[0].width = Cm(3.5)
                cells[1].width = Cm(11.5)
                cells[0].text = label
                cells[1].text = val
                self._set_cell_shading(cells[0], LIGHT_FILL_HEX)
                for c in cells: self._set_cell_margins(c)

            # 源码片段预览
            preview = wp.get('source_text_preview')
            if preview:
                note = doc.add_paragraph(style='Pack Note')
                note.add_run('原文背景：').bold = True
                note.add_run(preview)

    def add_figure_gallery(self, figures: List[Dict[str, Any]]):
        """添加图集区块"""
        doc = self.doc
        doc.add_page_break()
        doc.add_paragraph('图表证据包', style='Pack Heading 1')
        
        for fig in figures:
            p = doc.add_paragraph(style='Pack Heading 2')
            num = fig.get('figure_number', '')
            p.add_run(f"Figure {num}" if num else "Figure")
            self._add_bookmark(p, f"fig_{fig.get('figure_id', '')}")

            # 尝试插入图片
            img_item = fig.get('primary_single_figure', {}).get('page_crop_image') or fig.get('primary_single_figure', {}).get('raw_embedded_image')
            img_path = img_item.get('image_path') if img_item else None
            
            # Fallback check for alternative paths if direct path missing
            if not img_path:
                for candidate in fig.get('page_image_candidates', []):
                    if candidate.get('image_path'):
                        img_path = candidate.get('image_path')
                        break

            if img_path and Path(img_path).exists():
                try:
                    safe_path = self._normalize_image(Path(img_path))
                    width = self._get_image_width(safe_path)
                    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.paragraphs[-1].add_run().add_picture(str(safe_path), width=Inches(width))
                except Exception as e:
                    doc.add_paragraph(f"[图片载入失败: {e}]", style='Pack Note')

            # 图题
            cap = doc.add_paragraph(style='Pack Caption')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(fig.get('caption_prefix', f"Fig {num}. ") if num else "Fig. ")
            cap.add_run(fig.get('caption', ''))

            # 支撑逻辑说明
            if fig.get('supporting_claims'):
                note = doc.add_paragraph(style='Pack Note')
                note.add_run('支撑结论：').bold = True
                note.add_run(" | ".join(fig['supporting_claims'][:3]))

    def save(self):
        """保存文档"""
        self.doc.save(str(self.output_path))
        logger.info(f"Word document saved to {self.output_path}")

def generate_docx_report(material_pack_path: str | Path, output_docx: str | Path):
    """
    便捷函数：从写作材料包生成交付文档
    """
    path = Path(material_pack_path)
    if not path.exists():
        raise FileNotFoundError(f"Material pack not found at {path}")
    
    with open(path, 'r', encoding='utf-8') as f:
        pack = json.load(f)
        
    writer = WordWriter(output_docx)
    writer.setup()
    
    # 标题数据
    title = pack.get('paper_title', pack.get('goal', '文献分析报告'))
    metadata = {
        "分析目标": pack.get('goal', '未指定'),
        "产出时间": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "来源文件": Path(pack.get('source_pdf', 'unknown')).name,
        "数据契约": pack.get('schema_version', 'v2-standard')
    }
    
    writer.add_title_page(title, "期刊级文献多模态证据聚合交付稿", metadata)
    
    # 4. [v40.4] 期刊风格主题论述 (单栏紧凑版)
    themes = pack.get('semantic_themes', [])
    if themes:
        for i, theme in enumerate(themes, 1):
            h = writer.doc.add_heading(f"{i}. {theme['theme_title']}", level=1)
            h.paragraph_format.space_before = Pt(12)
            
            # 合成摘要
            p = writer.doc.add_paragraph(theme['summary'])
            p.paragraph_format.space_after = Pt(0) # 紧贴后续图片
            
            # 插入关联图片
            fig_ids = theme.get('linked_figure_ids', [])
            all_figures = pack.get('single_figure_cards', [])
            theme_figs = [f for f in all_figures if f['figure_id'] in fig_ids]
            if theme_figs:
                writer._add_journal_visuals(theme_figs)

            # 插入关联表格
            tab_ids = theme.get('linked_table_ids', [])
            all_tables = pack.get('single_table_cards', [])
            theme_tabs = [t for t in all_tables if t['table_id'] in tab_ids]
            if theme_tabs:
                writer._add_journal_visuals(theme_tabs, is_table=True)

    # 5. Appendix (原始证据链)
    writer.doc.add_heading("Appendix: Original Evidence Trace", level=1)
    points = pack.get('writing_point_cards', [])
    if points:
        writer.add_writing_points(points)
        
    writer.save()
    return str(Path(output_docx).resolve())

if __name__ == "__main__":
    # 简易测试入口
    import sys
    if len(sys.argv) > 2:
        generate_docx_report(sys.argv[1], sys.argv[2])
