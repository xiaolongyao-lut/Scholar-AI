# -*- coding: utf-8 -*-
"""Export API Router — TipTap content → formatted DOCX."""

from __future__ import annotations

import logging
import json
import os
import tempfile
import re
import shutil
import uuid
from dataclasses import dataclass, field, replace
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Literal

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field
from starlette.background import BackgroundTask

logger = logging.getLogger("ExportRouter")

router = APIRouter(prefix="/api/export", tags=["Export"])


def _safe_docx_filename_stem(value: str, fallback: str = "export") -> str:
    """Return a bounded filename stem for generated DOCX downloads.

    Args:
        value: User-facing title text, not a filesystem path.
        fallback: ASCII stem used when the title has no safe characters.

    Returns:
        A Windows-safe filename stem without path separators or control chars.
    """
    normalized = re.sub(r"\s+", " ", str(value or "").strip())
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", normalized).strip(" ._")
    if not safe:
        safe = fallback
    return safe[:96]


def _cleanup_export_tmp_dir(tmp_dir: Path) -> None:
    """Remove a generated DOCX temp directory after the response is sent.

    Args:
        tmp_dir: Directory returned by ``tempfile.mkdtemp(prefix="export_docx_")``.

    Returns:
        None. Paths outside the expected temp root are left untouched.
    """

    if not isinstance(tmp_dir, Path):
        raise TypeError("tmp_dir must be a pathlib.Path")

    resolved = tmp_dir.resolve()
    temp_root = Path(tempfile.gettempdir()).resolve()
    # First guard: parent must be temp_root and name must match expected prefix
    if resolved.parent != temp_root or not resolved.name.startswith("export_docx_"):
        return
    # Second guard: prevent directory traversal via symlinks or race conditions
    if not resolved.is_relative_to(temp_root):
        return
    if resolved.is_dir():
        shutil.rmtree(resolved)


class ExportDocxRequest(BaseModel):
    html: str = Field(..., min_length=1, max_length=500000)
    json_content: dict[str, Any] | None = Field(None, alias="json")
    title: str = Field("Untitled", max_length=200)
    style_profile: str | None = Field(default="gb_t_7714_review", max_length=80)
    verify_with_word: bool = False
    project_id: str | None = Field(default=None, max_length=128)
    require_action_preflight: bool = Field(
        default=False,
        description="When true, block DOCX export unless project workflow readiness allows export.",
    )


def _build_docx_action_preflight(req: ExportDocxRequest) -> dict[str, Any] | None | JSONResponse:
    """Return optional workflow preflight for project-scoped DOCX export.

    Args:
        req: Validated DOCX export request.

    Returns:
        Action preflight payload when a project id is available, otherwise None.

    Raises:
        HTTPException: If hard preflight is requested but no project id exists
            or the runtime gate blocks export readiness.
    """

    project_id = str(req.project_id or "").strip()
    if not project_id:
        if req.require_action_preflight:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "action_preflight_requires_project_id",
                    "message": "project_id is required when require_action_preflight is true.",
                },
            )
        return None
    try:
        from writing_runtime import get_writing_runtime

        runtime = get_writing_runtime()
        preflight = runtime.build_action_preflight(
            action_id="export.docx",
            required_claim_id="export_readiness",
            project_id=project_id,
            require_ready=bool(req.require_action_preflight),
            limit=500,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    if req.require_action_preflight and not bool(preflight.get("can_proceed")):
        return JSONResponse(
            status_code=409,
            content={
                "error": "action_preflight_blocked",
                "message": "Workflow passport and evidence integrity gate do not allow DOCX export yet.",
                "action_preflight": preflight,
            },
        )
    return preflight


def _preflight_header_value(preflight: dict[str, Any] | None) -> str:
    """Return a compact ASCII JSON header value for optional action preflight."""

    if preflight is None:
        return json.dumps(
            {
                "schema_version": "scholar_ai_action_preflight_v1",
                "status": "not_requested",
                "can_proceed": True,
            },
            ensure_ascii=True,
            sort_keys=True,
        )
    header = {
        "schema_version": preflight.get("schema_version"),
        "action_id": preflight.get("action_id"),
        "required_claim_id": preflight.get("required_claim_id"),
        "status": preflight.get("status"),
        "can_proceed": bool(preflight.get("can_proceed")),
        "claim_status": preflight.get("claim_status"),
        "gate_status": preflight.get("gate_status"),
        "refresh_required": bool(preflight.get("refresh_required")),
        "freshness_status": (preflight.get("freshness") or {}).get("status")
        if isinstance(preflight.get("freshness"), dict)
        else None,
    }
    return json.dumps(header, ensure_ascii=True, sort_keys=True)


class JournalStyleSpecDraftRequest(BaseModel):
    """Text-based official journal requirements submitted for profile drafting."""

    project_id: str = Field(..., min_length=1, max_length=128)
    journal_name: str = Field(..., min_length=1, max_length=160)
    spec_text: str = Field(..., min_length=20, max_length=120000)


class JournalStyleSpecConfirmRequest(BaseModel):
    """Confirm a project-scoped style profile draft for future exports."""

    project_id: str = Field(..., min_length=1, max_length=128)
    draft_id: str = Field(..., min_length=1, max_length=128)
    confirmed_by: str = Field("user", max_length=120)


@dataclass(frozen=True)
class _DocxRun:
    """Inline text plus formatting flags parsed from TipTap HTML."""

    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    superscript: bool = False
    subscript: bool = False
    formula_text: str | None = None
    equation_number: str | None = None


@dataclass
class _DocxBlock:
    """Block-level DOCX render unit produced by the HTML parser."""

    kind: Literal["paragraph", "heading", "list_item", "caption", "table"]
    runs: list[_DocxRun] = field(default_factory=list)
    heading_level: int | None = None
    table_rows: list[list[list[_DocxRun]]] = field(default_factory=list)
    header_rows: int = 0
    caption_kind: Literal["figure", "table"] | None = None


@dataclass
class _ExportQuality:
    """Small machine-readable export quality report for response headers."""

    citation_count: int = 0
    table_count: int = 0
    caption_count: int = 0
    crossref_count: int = 0
    formula_count: int = 0
    style_profile: str = "gb_t_7714_review"
    citation_style: str = "numeric"
    word_verify_status: Literal["skipped", "requested_unavailable"] = "skipped"

    def to_header(self) -> str:
        """Return an ASCII header payload safe for HTTP transport."""

        return (
            f"citations={self.citation_count};"
            f"tables={self.table_count};"
            f"captions={self.caption_count};"
            f"style_profile={self.style_profile};"
            f"citation_style={self.citation_style};"
            f"crossrefs={self.crossref_count};"
            f"formulas={self.formula_count};"
            f"word_verify={self.word_verify_status}"
        )


@dataclass(frozen=True)
class _StyleProfile:
    """Resolved journal/export style profile with bounded layout knobs."""

    profile_id: str
    citation_style: Literal["numeric", "author_year"]
    cjk_font: str
    heading_cjk_font: str
    latin_font: str
    body_pt: float
    title_pt: float
    top_margin_cm: float
    bottom_margin_cm: float
    left_margin_cm: float
    right_margin_cm: float


@dataclass
class _InlineRenderStats:
    """Counts emitted by rendering a bounded list of inline runs."""

    citation_count: int = 0
    crossref_count: int = 0
    formula_count: int = 0


@dataclass
class _DocxRenderContext:
    """Mutable render state for stable Word bookmark and caption numbering."""

    caption_counters: dict[str, int] = field(default_factory=lambda: {"figure": 0, "table": 0})
    formula_counter: int = 0
    bookmark_next_id: int = 1

    def next_bookmark_id(self) -> int:
        """Return a document-local bookmark id suitable for paired markers."""

        value = self.bookmark_next_id
        self.bookmark_next_id += 1
        return value


_STYLE_PROFILES: dict[str, _StyleProfile] = {
    "gb_t_7714_review": _StyleProfile(
        profile_id="gb_t_7714_review",
        citation_style="numeric",
        cjk_font="宋体",
        heading_cjk_font="黑体",
        latin_font="Times New Roman",
        body_pt=10.5,
        title_pt=18.0,
        top_margin_cm=2.2,
        bottom_margin_cm=2.0,
        left_margin_cm=2.2,
        right_margin_cm=2.2,
    ),
    "ieee": _StyleProfile(
        profile_id="ieee",
        citation_style="numeric",
        cjk_font="宋体",
        heading_cjk_font="黑体",
        latin_font="Times New Roman",
        body_pt=10.0,
        title_pt=16.0,
        top_margin_cm=1.9,
        bottom_margin_cm=1.9,
        left_margin_cm=1.8,
        right_margin_cm=1.8,
    ),
    "apa": _StyleProfile(
        profile_id="apa",
        citation_style="author_year",
        cjk_font="宋体",
        heading_cjk_font="黑体",
        latin_font="Times New Roman",
        body_pt=12.0,
        title_pt=16.0,
        top_margin_cm=2.54,
        bottom_margin_cm=2.54,
        left_margin_cm=2.54,
        right_margin_cm=2.54,
    ),
    "nature": _StyleProfile(
        profile_id="nature",
        citation_style="numeric",
        cjk_font="宋体",
        heading_cjk_font="黑体",
        latin_font="Arial",
        body_pt=10.0,
        title_pt=17.0,
        top_margin_cm=2.0,
        bottom_margin_cm=2.0,
        left_margin_cm=2.0,
        right_margin_cm=2.0,
    ),
    "generic_academic": _StyleProfile(
        profile_id="generic_academic",
        citation_style="numeric",
        cjk_font="宋体",
        heading_cjk_font="黑体",
        latin_font="Times New Roman",
        body_pt=11.0,
        title_pt=18.0,
        top_margin_cm=2.2,
        bottom_margin_cm=2.2,
        left_margin_cm=2.2,
        right_margin_cm=2.2,
    ),
}

_JOURNAL_SPEC_ALLOWED_SUFFIXES = {".txt", ".md", ".markdown"}
_JOURNAL_SPEC_MAX_UPLOAD_BYTES = 256_000
_JOURNAL_SPEC_MAX_TEXT_CHARS = 120_000
_JOURNAL_SPEC_ALLOWED_MEDIA_TYPES = {
    "text/plain",
    "text/markdown",
    "application/octet-stream",
}


_ACADEMIC_CITATION_RE = re.compile(
    r"(\[(?:\d+(?:\s*[-,，]\s*\d+)*|chunk:[^\]]{1,200}|evidence_pack:[^\]]{1,200})\]"
    r"|\[\^cite:[^\]]{1,200}\])"
)
_CAPTION_RE = re.compile(r"^\s*(图|表|Figure|Table)\s*(?:\d+|[一二三四五六七八九十]+)?[\s.．:：、-]*(.*)$", re.I)
_CROSS_REFERENCE_RE = re.compile(
    r"(?P<figure>(?:图|Figure)\s*(?P<figure_number>\d{1,3}))"
    r"|(?P<table>(?:表|Table)\s*(?P<table_number>\d{1,3}))"
    r"|(?P<equation>(?:式|Equation|Eq\.?)\s*[（(]?\s*(?P<equation_number>\d{1,3})\s*[）)]?)",
    re.I,
)
_FORMULA_LABEL_RE = re.compile(r"^\s*(?:式|Equation|Eq\.?)\s*[（(]?\s*\d{1,3}\s*[）)]?\s*[:：]?\s*$", re.I)


def _utc_now() -> str:
    """Return an ISO timestamp for local audit metadata."""

    return datetime.now(timezone.utc).isoformat()


def _safe_profile_token(value: str, *, fallback: str = "journal") -> str:
    """Return a bounded lowercase token used in custom profile ids."""

    normalized = str(value or "").strip().lower().replace("-", "_")
    normalized = re.sub(r"[^a-z0-9_]+", "_", normalized)
    normalized = re.sub(r"_+", "_", normalized).strip("_")
    return (normalized or fallback)[:48]


def _project_style_profile_path(project_id: str) -> Path:
    """Return the project-scoped journal style profile store path."""

    from project_paths import project_data_path

    normalized_project_id = str(project_id or "").strip()
    if not normalized_project_id:
        raise ValueError("project_id is required")
    return project_data_path(normalized_project_id, "journal_style_profiles.json")


def _atomic_write_json(path: Path, payload: dict[str, Any]) -> None:
    """Write JSON atomically so confirmed profiles are not partially persisted."""

    path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = path.with_name(f"{path.name}.{uuid.uuid4().hex}.tmp")
    tmp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
    os.replace(tmp_path, path)


def _load_project_style_profile_store(project_id: str) -> dict[str, Any]:
    """Load project-scoped journal profile drafts and confirmed profiles."""

    path = _project_style_profile_path(project_id)
    if not path.exists():
        return {"drafts": {}, "profiles": {}}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"drafts": {}, "profiles": {}}
    if not isinstance(payload, dict):
        return {"drafts": {}, "profiles": {}}
    drafts = payload.get("drafts")
    profiles = payload.get("profiles")
    return {
        "drafts": drafts if isinstance(drafts, dict) else {},
        "profiles": profiles if isinstance(profiles, dict) else {},
    }


def _save_project_style_profile_store(project_id: str, payload: dict[str, Any]) -> None:
    """Persist project style profiles under workspace artifacts."""

    _atomic_write_json(_project_style_profile_path(project_id), payload)


def _bounded_spec_text(value: str) -> str:
    """Normalize official requirement text before deterministic extraction."""

    text = str(value or "").replace("\x00", "")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) < 20:
        raise ValueError("spec_text must contain at least 20 characters")
    if len(text) > _JOURNAL_SPEC_MAX_TEXT_CHARS:
        raise ValueError(f"spec_text must be at most {_JOURNAL_SPEC_MAX_TEXT_CHARS} characters")
    return text


def _extract_body_pt(text: str, default: float) -> float:
    """Extract a bounded body font size from journal instructions."""

    match = re.search(r"(?P<size>8|9|10|10\.5|11|12|13|14)\s*(?:pt|point|points|磅|号)", text, re.I)
    if not match:
        return default
    return float(match.group("size"))


def _extract_margin_cm(text: str, default: float) -> float:
    """Extract a bounded margin value in centimeters."""

    cm_match = re.search(r"(?P<size>1(?:\.\d+)?|2(?:\.\d+)?|3(?:\.\d+)?)\s*(?:cm|厘米)", text, re.I)
    if cm_match:
        return max(1.0, min(3.5, float(cm_match.group("size"))))
    inch_match = re.search(r"(?P<size>0\.5|0\.75|1(?:\.0)?|1\.25)\s*(?:in|inch|inches)", text, re.I)
    if inch_match:
        return round(max(1.0, min(3.5, float(inch_match.group("size")) * 2.54)), 2)
    return default


def _extract_latin_font(text: str, default: str) -> str:
    """Extract a safe known Latin font from style instructions."""

    lowered = text.lower()
    if "arial" in lowered:
        return "Arial"
    if "calibri" in lowered:
        return "Calibri"
    if "times new roman" in lowered or "times" in lowered:
        return "Times New Roman"
    return default


def _extract_citation_style(text: str, default: Literal["numeric", "author_year"]) -> Literal["numeric", "author_year"]:
    """Extract bounded citation style from journal instructions."""

    lowered = text.lower()
    if "apa" in lowered or "author-year" in lowered or "author year" in lowered or "作者" in text:
        return "author_year"
    if "ieee" in lowered or "numeric" in lowered or "numbered" in lowered or "顺序编码" in text:
        return "numeric"
    return default


def _profile_to_public(profile: dict[str, Any]) -> dict[str, Any]:
    """Return a compact public journal profile payload."""

    return {
        "profile_id": profile["profile_id"],
        "journal_name": profile["journal_name"],
        "citation_style": profile["citation_style"],
        "latin_font": profile["latin_font"],
        "cjk_font": profile["cjk_font"],
        "heading_cjk_font": profile["heading_cjk_font"],
        "body_pt": profile["body_pt"],
        "title_pt": profile["title_pt"],
        "margins_cm": {
            "top": profile["top_margin_cm"],
            "bottom": profile["bottom_margin_cm"],
            "left": profile["left_margin_cm"],
            "right": profile["right_margin_cm"],
        },
        "figure_caption_position": profile["figure_caption_position"],
        "table_caption_position": profile["table_caption_position"],
    }


def _draft_journal_style_profile(
    *,
    project_id: str,
    journal_name: str,
    spec_text: str,
    source: dict[str, Any],
) -> dict[str, Any]:
    """Create a reviewable project-scoped style profile draft."""

    if not project_id.strip():
        raise ValueError("project_id is required")
    text = _bounded_spec_text(spec_text)
    base = _STYLE_PROFILES["generic_academic"]
    citation_style = _extract_citation_style(text, base.citation_style)
    margin = _extract_margin_cm(text, base.top_margin_cm)
    body_pt = _extract_body_pt(text, base.body_pt)
    latin_font = _extract_latin_font(text, base.latin_font)
    journal_token = _safe_profile_token(journal_name)
    draft_id = f"style_draft_{uuid.uuid4().hex[:12]}"
    profile_id = f"custom_{journal_token}_{uuid.uuid4().hex[:8]}"
    warnings: list[str] = []
    if citation_style == "author_year":
        warnings.append("Detected APA or author-year citation requirements; confirm reference-list details manually.")
    else:
        warnings.append("Detected numeric citation requirements; confirm exact bibliography punctuation manually.")
    if "figure" in text.lower() or "图" in text:
        warnings.append("Detected figure-caption instructions; verify figure caption placement before final export.")
    if "table" in text.lower() or "表" in text:
        warnings.append("Detected table-caption instructions; verify table caption placement before final export.")

    profile = {
        "profile_id": profile_id,
        "journal_name": str(journal_name or "").strip()[:160],
        "citation_style": citation_style,
        "cjk_font": base.cjk_font,
        "heading_cjk_font": base.heading_cjk_font,
        "latin_font": latin_font,
        "body_pt": body_pt,
        "title_pt": max(body_pt + 4.0, 16.0),
        "top_margin_cm": margin,
        "bottom_margin_cm": margin,
        "left_margin_cm": margin,
        "right_margin_cm": margin,
        "figure_caption_position": "below",
        "table_caption_position": "above" if re.search(r"table captions?.{0,40}above|表题.{0,20}上", text, re.I) else "below",
        "created_at": _utc_now(),
        "source": source,
    }
    return {
        "draft_id": draft_id,
        "project_id": project_id,
        "status": "draft",
        "profile": profile,
        "warnings": warnings,
        "requires_confirmation": True,
        "created_at": profile["created_at"],
        "source": source,
    }


def _confirmed_profile_to_style(profile: dict[str, Any]) -> _StyleProfile:
    """Convert a confirmed project profile record into internal style knobs."""

    profile_id = str(profile.get("profile_id") or "").strip()
    if not profile_id.startswith("custom_"):
        raise ValueError("custom style profile id is invalid")
    citation_style = str(profile.get("citation_style") or "numeric")
    if citation_style not in {"numeric", "author_year"}:
        citation_style = "numeric"
    return _StyleProfile(
        profile_id=profile_id,
        citation_style=citation_style,  # type: ignore[arg-type]
        cjk_font=str(profile.get("cjk_font") or "宋体")[:80],
        heading_cjk_font=str(profile.get("heading_cjk_font") or "黑体")[:80],
        latin_font=str(profile.get("latin_font") or "Times New Roman")[:80],
        body_pt=float(profile.get("body_pt") or 11.0),
        title_pt=float(profile.get("title_pt") or 18.0),
        top_margin_cm=float(profile.get("top_margin_cm") or 2.2),
        bottom_margin_cm=float(profile.get("bottom_margin_cm") or 2.2),
        left_margin_cm=float(profile.get("left_margin_cm") or 2.2),
        right_margin_cm=float(profile.get("right_margin_cm") or 2.2),
    )


def _resolve_style_profile(value: str | None, project_id: str | None = None) -> _StyleProfile:
    """Resolve a supported journal/export style profile.

    Args:
        value: Optional profile id. Hyphens are normalized to underscores.
        project_id: Required when resolving project-scoped custom profiles.

    Returns:
        Bounded style profile used for DOCX layout and quality reporting.
    """

    raw_value = str(value or "gb_t_7714_review").strip().lower().replace("-", "_")
    if not raw_value:
        raw_value = "gb_t_7714_review"
    try:
        return _STYLE_PROFILES[raw_value]
    except KeyError as exc:
        if raw_value.startswith("custom_") and isinstance(project_id, str) and project_id.strip():
            store = _load_project_style_profile_store(project_id)
            profile = store["profiles"].get(raw_value)
            if isinstance(profile, dict):
                return _confirmed_profile_to_style(profile)
        allowed = ", ".join(sorted(_STYLE_PROFILES))
        raise ValueError(f"unsupported style_profile: {raw_value}; allowed: {allowed}") from exc


def _runs_text(runs: list[_DocxRun]) -> str:
    """Return plain text for a run list."""

    return "".join(run.text for run in runs)


def _trim_runs(runs: list[_DocxRun]) -> list[_DocxRun]:
    """Trim only block-edge whitespace while preserving inline spacing."""

    cleaned = [run for run in runs if run.text]
    if not cleaned:
        return []
    first = replace(cleaned[0], text=cleaned[0].text.lstrip())
    last = replace(cleaned[-1], text=cleaned[-1].text.rstrip()) if len(cleaned) > 1 else first
    if len(cleaned) == 1:
        cleaned = [first]
    else:
        cleaned = [first, *cleaned[1:-1], last]
    return [run for run in cleaned if run.text]


def _merge_run(target: list[_DocxRun], run: _DocxRun) -> None:
    """Append a run, merging identical adjacent formatting to reduce DOCX noise."""

    if not run.text:
        return
    if run.formula_text:
        target.append(run)
        return
    if target:
        previous = target[-1]
        if previous.formula_text:
            target.append(run)
            return
        if (
            previous.bold == run.bold
            and previous.italic == run.italic
            and previous.underline == run.underline
            and previous.superscript == run.superscript
            and previous.subscript == run.subscript
            and previous.formula_text == run.formula_text
            and previous.equation_number == run.equation_number
        ):
            target[-1] = replace(previous, text=f"{previous.text}{run.text}")
            return
    target.append(run)


class _TipTapDocxParser(HTMLParser):
    """Parse the TipTap HTML subset needed for academic DOCX export."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[_DocxBlock] = []
        self._runs: list[_DocxRun] = []
        self._kind: Literal["paragraph", "heading", "list_item", "caption"] | None = None
        self._heading_level: int | None = None
        self._list_depth = 0
        self._bold_depth = 0
        self._italic_depth = 0
        self._underline_depth = 0
        self._superscript_depth = 0
        self._subscript_depth = 0
        self._span_stack: list[tuple[bool, bool]] = []
        self._table_rows: list[list[list[_DocxRun]]] | None = None
        self._table_row: list[list[_DocxRun]] | None = None
        self._cell_runs: list[_DocxRun] | None = None
        self._current_row_has_header = False
        self._header_rows = 0
        self._formula_span_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        """Handle an HTML opening tag from the TipTap subset."""

        normalized = tag.lower()
        attr_map = {key.lower(): value or "" for key, value in attrs}
        if self._formula_span_depth > 0:
            if normalized == "span":
                self._formula_span_depth += 1
            return
        if normalized in {"strong", "b"}:
            self._bold_depth += 1
            return
        if normalized in {"em", "i"}:
            self._italic_depth += 1
            return
        if normalized == "u":
            self._underline_depth += 1
            return
        if normalized == "sup":
            self._superscript_depth += 1
            return
        if normalized == "sub":
            self._subscript_depth += 1
            return
        if normalized == "span":
            formula = attr_map.get("data-formula", "").strip()
            if formula:
                self._append_formula(formula, attr_map.get("data-equation-number", ""))
                self._formula_span_depth += 1
                return
            style = attr_map.get("style", "").lower()
            is_super = "vertical-align: super" in style or "baseline-shift: super" in style
            is_sub = "vertical-align: sub" in style or "baseline-shift: sub" in style
            self._span_stack.append((is_super, is_sub))
            self._superscript_depth += int(is_super)
            self._subscript_depth += int(is_sub)
            return
        if normalized == "br":
            self._append_text("\n")
            return
        if normalized in {"h1", "h2", "h3", "h4"} and self._table_rows is None:
            self._flush_text_block()
            self._kind = "heading"
            self._heading_level = min(int(normalized[1]), 3)
            return
        if normalized in {"p", "div", "blockquote"} and self._table_rows is None:
            self._flush_text_block()
            self._kind = "paragraph"
            return
        if normalized in {"ul", "ol"} and self._table_rows is None:
            self._list_depth += 1
            return
        if normalized == "li" and self._table_rows is None:
            self._flush_text_block()
            self._kind = "list_item"
            return
        if normalized == "figcaption" and self._table_rows is None:
            self._flush_text_block()
            self._kind = "caption"
            return
        if normalized == "table":
            self._flush_text_block()
            self._table_rows = []
            self._table_row = None
            self._cell_runs = None
            self._current_row_has_header = False
            self._header_rows = 0
            return
        if normalized == "tr" and self._table_rows is not None:
            self._table_row = []
            self._current_row_has_header = False
            return
        if normalized in {"td", "th"} and self._table_rows is not None:
            self._cell_runs = []
            if normalized == "th":
                self._current_row_has_header = True
                self._bold_depth += 1

    def handle_endtag(self, tag: str) -> None:
        """Handle an HTML closing tag from the TipTap subset."""

        normalized = tag.lower()
        if self._formula_span_depth > 0:
            if normalized == "span":
                self._formula_span_depth = max(0, self._formula_span_depth - 1)
            return
        if normalized in {"strong", "b"}:
            self._bold_depth = max(0, self._bold_depth - 1)
            return
        if normalized in {"em", "i"}:
            self._italic_depth = max(0, self._italic_depth - 1)
            return
        if normalized == "u":
            self._underline_depth = max(0, self._underline_depth - 1)
            return
        if normalized == "sup":
            self._superscript_depth = max(0, self._superscript_depth - 1)
            return
        if normalized == "sub":
            self._subscript_depth = max(0, self._subscript_depth - 1)
            return
        if normalized == "span":
            if self._span_stack:
                was_super, was_sub = self._span_stack.pop()
                self._superscript_depth = max(0, self._superscript_depth - int(was_super))
                self._subscript_depth = max(0, self._subscript_depth - int(was_sub))
            return
        if normalized in {"h1", "h2", "h3", "h4", "p", "div", "blockquote", "li", "figcaption"}:
            if self._table_rows is None:
                self._flush_text_block()
            return
        if normalized in {"ul", "ol"} and self._table_rows is None:
            self._list_depth = max(0, self._list_depth - 1)
            return
        if normalized in {"td", "th"} and self._table_rows is not None:
            if self._table_row is not None:
                self._table_row.append(_trim_runs(self._cell_runs or []))
            self._cell_runs = None
            if normalized == "th":
                self._bold_depth = max(0, self._bold_depth - 1)
            return
        if normalized == "tr" and self._table_rows is not None:
            if self._table_row:
                self._table_rows.append(self._table_row)
                if self._current_row_has_header:
                    self._header_rows += 1
            self._table_row = None
            self._current_row_has_header = False
            return
        if normalized == "table" and self._table_rows is not None:
            rows = [row for row in self._table_rows if row]
            if rows:
                self.blocks.append(
                    _DocxBlock(
                        kind="table",
                        table_rows=rows,
                        header_rows=self._header_rows,
                    )
                )
            self._table_rows = None
            self._table_row = None
            self._cell_runs = None
            self._header_rows = 0

    def handle_data(self, data: str) -> None:
        """Append text data into the active block or table cell."""

        if self._formula_span_depth > 0:
            return
        self._append_text(data)

    def close(self) -> None:
        """Flush any unterminated block before closing the parser."""

        self._flush_text_block()
        if self._table_rows is not None:
            rows = [row for row in self._table_rows if row]
            if rows:
                self.blocks.append(_DocxBlock(kind="table", table_rows=rows, header_rows=self._header_rows))
        super().close()

    def _target_runs(self) -> list[_DocxRun]:
        """Return the active run buffer."""

        if self._table_rows is not None and self._cell_runs is not None:
            return self._cell_runs
        if self._kind is None:
            self._kind = "paragraph"
        return self._runs

    def _append_text(self, data: str) -> None:
        """Append normalized text using the current inline formatting state."""

        if not data:
            return
        text = re.sub(r"[ \t\r\n\f\v]+", " ", data)
        target = self._target_runs()
        if not text.strip():
            if target and not target[-1].text.endswith((" ", "\n")):
                _merge_run(target, replace(target[-1], text=" "))
            return
        if not target:
            text = text.lstrip()
        run = _DocxRun(
            text=text,
            bold=self._bold_depth > 0,
            italic=self._italic_depth > 0,
            underline=self._underline_depth > 0,
            superscript=self._superscript_depth > 0,
            subscript=self._subscript_depth > 0,
        )
        _merge_run(target, run)

    def _append_formula(self, formula: str, equation_number: str) -> None:
        """Append a formula placeholder whose target shape is stored in attrs."""

        text = re.sub(r"\s+", " ", formula).strip()
        if not text:
            return
        number_match = re.search(r"\d{1,3}", str(equation_number or ""))
        target = self._target_runs()
        _merge_run(
            target,
            _DocxRun(
                text=text,
                formula_text=text,
                equation_number=number_match.group(0) if number_match else None,
            ),
        )

    def _flush_text_block(self) -> None:
        """Append the current text block when it has non-empty content."""

        runs = _trim_runs(self._runs)
        if not runs:
            self._runs = []
            self._kind = None
            self._heading_level = None
            return
        kind = self._kind or ("list_item" if self._list_depth else "paragraph")
        caption_kind = _infer_caption_kind(_runs_text(runs)) if kind == "caption" else None
        self.blocks.append(
            _DocxBlock(
                kind=kind,
                runs=runs,
                heading_level=self._heading_level if kind == "heading" else None,
                caption_kind=caption_kind,
            )
        )
        self._runs = []
        self._kind = None
        self._heading_level = None


def _infer_caption_kind(text: str) -> Literal["figure", "table"] | None:
    """Infer figure/table caption type from localized caption text."""

    match = _CAPTION_RE.match(text)
    if not match:
        return None
    label = match.group(1).lower()
    return "table" if label in {"表", "table"} else "figure"


def _split_citation_runs(run: _DocxRun) -> list[_DocxRun]:
    """Split citation tokens so they can be rendered as superscript runs."""

    if run.formula_text or run.superscript or run.subscript:
        return [run]
    parts: list[_DocxRun] = []
    last_index = 0
    for match in _ACADEMIC_CITATION_RE.finditer(run.text):
        if match.start() > last_index:
            parts.append(replace(run, text=run.text[last_index:match.start()]))
        parts.append(replace(run, text=match.group(0), superscript=True))
        last_index = match.end()
    if last_index < len(run.text):
        parts.append(replace(run, text=run.text[last_index:]))
    return [part for part in parts if part.text]


def _set_run_fonts(run: Any, *, profile: _StyleProfile, heading: bool = False) -> None:
    """Apply the repository's CJK/Latin font convention to one DOCX run."""

    from docx.oxml.ns import qn

    run.font.name = profile.latin_font
    run._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), profile.heading_cjk_font if heading else profile.cjk_font)


def _append_plain_run(
    paragraph: Any,
    text: str,
    run_info: _DocxRun,
    *,
    profile: _StyleProfile,
    heading: bool = False,
) -> None:
    """Append one formatted text run when no special field handling is needed."""

    if not text:
        return
    run = paragraph.add_run(text)
    _set_run_fonts(run, profile=profile, heading=heading)
    run.bold = run_info.bold
    run.italic = run_info.italic
    run.underline = run_info.underline
    run.font.superscript = run_info.superscript
    run.font.subscript = run_info.subscript


def _append_bookmark_start(paragraph: Any, name: str, bookmark_id: int) -> None:
    """Append a Word bookmark start marker for later REF fields."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement("w:bookmarkStart")
    element.set(qn("w:id"), str(bookmark_id))
    element.set(qn("w:name"), name)
    paragraph._p.append(element)


def _append_bookmark_end(paragraph: Any, bookmark_id: int) -> None:
    """Append a Word bookmark end marker paired with a prior start marker."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement("w:bookmarkEnd")
    element.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(element)


def _append_ref_field(paragraph: Any, bookmark_name: str, display_text: str) -> None:
    """Append a Word REF field pointing at a generated bookmark name."""

    _append_field_run(paragraph, f" REF {bookmark_name} \\h ", display_text)


def _cross_reference_target(match: re.Match[str]) -> tuple[str, str, str, str] | None:
    """Return kind, number, prefix, and suffix for a body cross-reference."""

    text = match.group(0)
    if match.group("figure"):
        number = match.group("figure_number")
        return "figure", number, text[: text.rfind(number)], text[text.rfind(number) + len(number) :]
    if match.group("table"):
        number = match.group("table_number")
        return "table", number, text[: text.rfind(number)], text[text.rfind(number) + len(number) :]
    if match.group("equation"):
        number = match.group("equation_number")
        return "equation", number, text[: text.rfind(number)], text[text.rfind(number) + len(number) :]
    return None


def _add_text_run_with_crossrefs(
    paragraph: Any,
    run_info: _DocxRun,
    *,
    profile: _StyleProfile,
    heading: bool = False,
    suppress_equation_crossrefs: bool = False,
) -> int:
    """Append a text run while converting body figure/table/equation refs."""

    if run_info.superscript or run_info.subscript:
        _append_plain_run(paragraph, run_info.text, run_info, profile=profile, heading=heading)
        return 0

    suppress_all_equation_crossrefs = suppress_equation_crossrefs and bool(_FORMULA_LABEL_RE.fullmatch(run_info.text))
    crossref_count = 0
    last_index = 0
    for match in _CROSS_REFERENCE_RE.finditer(run_info.text):
        if match.start() > last_index:
            _append_plain_run(
                paragraph,
                run_info.text[last_index : match.start()],
                run_info,
                profile=profile,
                heading=heading,
            )
        target = _cross_reference_target(match)
        if target is None:
            _append_plain_run(paragraph, match.group(0), run_info, profile=profile, heading=heading)
        else:
            kind, number, prefix, suffix = target
            if kind == "equation" and suppress_all_equation_crossrefs:
                _append_plain_run(paragraph, match.group(0), run_info, profile=profile, heading=heading)
            else:
                _append_plain_run(paragraph, prefix, run_info, profile=profile, heading=heading)
                _append_ref_field(paragraph, f"litassist_{kind}_{number}", number)
                _append_plain_run(paragraph, suffix, run_info, profile=profile, heading=heading)
                crossref_count += 1
        last_index = match.end()

    if last_index < len(run_info.text):
        _append_plain_run(paragraph, run_info.text[last_index:], run_info, profile=profile, heading=heading)
    return crossref_count


def _append_omml_formula(
    paragraph: Any,
    formula_text: str,
    *,
    equation_number: str | None,
    context: _DocxRenderContext,
    profile: _StyleProfile,
) -> None:
    """Append a lightweight OMML formula with a bookmarkable equation number."""

    from docx.oxml import OxmlElement

    if not formula_text.strip():
        raise ValueError("formula_text must be non-empty")

    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    math_text.text = formula_text.strip()
    math_run.append(math_text)
    omath.append(math_run)
    paragraph._p.append(omath)

    number = equation_number or str(context.formula_counter + 1)
    context.formula_counter = max(context.formula_counter, int(number))
    bookmark_id = context.next_bookmark_id()
    _append_plain_run(
        paragraph,
        " ",
        _DocxRun(text=" "),
        profile=profile,
    )
    _append_bookmark_start(paragraph, f"litassist_equation_{number}", bookmark_id)
    _append_plain_run(
        paragraph,
        f"({number})",
        _DocxRun(text=f"({number})"),
        profile=profile,
    )
    _append_bookmark_end(paragraph, bookmark_id)


def _add_docx_runs_with_stats(
    paragraph: Any,
    runs: list[_DocxRun],
    *,
    profile: _StyleProfile,
    heading: bool = False,
    context: _DocxRenderContext | None = None,
) -> _InlineRenderStats:
    """Append formatted runs and return citation/crossref/formula counts."""

    stats = _InlineRenderStats()
    active_context = context or _DocxRenderContext()
    suppress_equation_crossrefs = any(run.formula_text for run in runs)
    for source_run in runs:
        for run_info in _split_citation_runs(source_run):
            if run_info.formula_text:
                _append_omml_formula(
                    paragraph,
                    run_info.formula_text,
                    equation_number=run_info.equation_number,
                    context=active_context,
                    profile=profile,
                )
                stats.formula_count += 1
                continue
            if run_info.superscript and _ACADEMIC_CITATION_RE.fullmatch(run_info.text):
                _append_plain_run(paragraph, run_info.text, run_info, profile=profile, heading=heading)
                stats.citation_count += 1
                continue
            stats.crossref_count += _add_text_run_with_crossrefs(
                paragraph,
                run_info,
                profile=profile,
                heading=heading,
                suppress_equation_crossrefs=suppress_equation_crossrefs,
            )
    return stats


def _append_field_run(paragraph: Any, instruction: str, display_text: str) -> None:
    """Append a simple Word field sequence used for figure/table captions."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_char)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(fld_char)

    paragraph.add_run(display_text)

    end = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    end._r.append(fld_char)


def _append_bookmarked_field_run(
    paragraph: Any,
    instruction: str,
    display_text: str,
    *,
    bookmark_name: str,
    context: _DocxRenderContext,
) -> None:
    """Append a Word field wrapped by a bookmark for cross-reference targets."""

    bookmark_id = context.next_bookmark_id()
    _append_bookmark_start(paragraph, bookmark_name, bookmark_id)
    _append_field_run(paragraph, instruction, display_text)
    _append_bookmark_end(paragraph, bookmark_id)


def _set_border_element(parent: Any, name: str, *, val: str, size: str = "8", color: str = "000000") -> None:
    """Append one WordprocessingML border element."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement(f"w:{name}")
    element.set(qn("w:val"), val)
    if val != "nil":
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
    parent.append(element)


def _apply_three_line_table_borders(table: Any, *, header_rows: int) -> None:
    """Apply a simple three-line academic table border model."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    _set_border_element(borders, "top", val="single", size="12")
    _set_border_element(borders, "left", val="nil")
    _set_border_element(borders, "bottom", val="single", size="12")
    _set_border_element(borders, "right", val="nil")
    _set_border_element(borders, "insideH", val="nil")
    _set_border_element(borders, "insideV", val="nil")
    tbl_pr.append(borders)

    for row in table.rows[: max(0, header_rows)]:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            existing_cell_borders = tc_pr.find(qn("w:tcBorders"))
            if existing_cell_borders is not None:
                tc_pr.remove(existing_cell_borders)
            cell_borders = OxmlElement("w:tcBorders")
            _set_border_element(cell_borders, "bottom", val="single", size="8")
            tc_pr.append(cell_borders)


def _add_table_block(
    document: Any,
    block: _DocxBlock,
    *,
    profile: _StyleProfile,
    context: _DocxRenderContext,
) -> _InlineRenderStats:
    """Append one parsed HTML table and return rendered inline counts."""

    rows = block.table_rows
    if not rows:
        return _InlineRenderStats()
    column_count = max(len(row) for row in rows)
    if column_count <= 0:
        return _InlineRenderStats()
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    stats = _InlineRenderStats()
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell_runs = row[column_index] if column_index < len(row) else []
            paragraph = cell.paragraphs[0]
            if row_index < block.header_rows:
                cell_runs = [replace(run, bold=True) for run in cell_runs]
            cell_stats = _add_docx_runs_with_stats(paragraph, cell_runs, profile=profile, context=context)
            stats.citation_count += cell_stats.citation_count
            stats.crossref_count += cell_stats.crossref_count
            stats.formula_count += cell_stats.formula_count
    _apply_three_line_table_borders(table, header_rows=block.header_rows)
    return stats


def _caption_text_without_label(text: str) -> tuple[Literal["figure", "table"], str, str]:
    """Return caption kind, localized label, and text after any static number."""

    match = _CAPTION_RE.match(text)
    if not match:
        return "figure", "图", text.strip()
    raw_label = match.group(1)
    kind: Literal["figure", "table"] = "table" if raw_label.lower() in {"表", "table"} else "figure"
    label = "表" if kind == "table" and raw_label == "表" else "Table" if kind == "table" else "图" if raw_label == "图" else "Figure"
    return kind, label, match.group(2).strip()


def _add_caption_block(document: Any, block: _DocxBlock, *, profile: _StyleProfile, context: _DocxRenderContext) -> None:
    """Append a figure/table caption with a Word SEQ field placeholder."""

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = _runs_text(block.runs)
    kind, label, suffix = _caption_text_without_label(text)
    context.caption_counters[kind] = context.caption_counters.get(kind, 0) + 1
    number = str(context.caption_counters[kind])
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label} ")
    _set_run_fonts(label_run, profile=profile, heading=True)
    label_run.bold = True
    _append_bookmarked_field_run(
        paragraph,
        f" SEQ {'Table' if kind == 'table' else 'Figure'} \\* ARABIC ",
        number,
        bookmark_name=f"litassist_{kind}_{number}",
        context=context,
    )
    if suffix:
        suffix_run = paragraph.add_run(f" {suffix}")
        _set_run_fonts(suffix_run, profile=profile)


def _render_blocks_to_docx(
    document: Any,
    blocks: list[_DocxBlock],
    quality: _ExportQuality,
    *,
    profile: _StyleProfile,
) -> None:
    """Render parsed blocks into a python-docx document."""

    heading_styles = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    context = _DocxRenderContext()
    for block in blocks:
        if block.kind == "table":
            quality.table_count += 1
            stats = _add_table_block(document, block, profile=profile, context=context)
            quality.citation_count += stats.citation_count
            quality.crossref_count += stats.crossref_count
            quality.formula_count += stats.formula_count
            continue
        if block.kind == "caption":
            quality.caption_count += 1
            _add_caption_block(document, block, profile=profile, context=context)
            continue
        if block.kind == "heading":
            style = heading_styles.get(block.heading_level or 1, "Heading 1")
            paragraph = document.add_paragraph(style=style)
            stats = _add_docx_runs_with_stats(paragraph, block.runs, profile=profile, heading=True, context=context)
            quality.citation_count += stats.citation_count
            quality.crossref_count += stats.crossref_count
            quality.formula_count += stats.formula_count
            continue
        paragraph = document.add_paragraph(style="List Bullet" if block.kind == "list_item" else "Normal")
        stats = _add_docx_runs_with_stats(paragraph, block.runs, profile=profile, context=context)
        quality.citation_count += stats.citation_count
        quality.crossref_count += stats.crossref_count
        quality.formula_count += stats.formula_count


def _html_to_docx(
    html: str,
    title: str,
    output_path: Path,
    style_profile: str | None = None,
    verify_with_word: bool = False,
    project_id: str | None = None,
) -> tuple[Path, _ExportQuality]:
    """Convert TipTap HTML to DOCX using existing WordWriter infrastructure."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")

    if output_path.suffix.lower() != ".docx":
        raise ValueError("output_path must use a .docx suffix")
    if not isinstance(html, str) or not html.strip():
        raise ValueError("html must be non-empty")

    doc = Document()
    profile = _resolve_style_profile(style_profile, project_id=project_id)
    quality = _ExportQuality(
        style_profile=profile.profile_id,
        citation_style=profile.citation_style,
        word_verify_status="requested_unavailable" if verify_with_word else "skipped",
    )

    # Page setup — reuse WordWriter conventions
    sec = doc.sections[0]
    sec.top_margin = Cm(profile.top_margin_cm)
    sec.bottom_margin = Cm(profile.bottom_margin_cm)
    sec.left_margin = Cm(profile.left_margin_cm)
    sec.right_margin = Cm(profile.right_margin_cm)

    # CJK/Latin dual font — reuse WordWriter pattern
    styles = doc.styles
    normal_font = styles["Normal"].font
    normal_font.name = profile.latin_font
    normal_font.size = Pt(profile.body_pt)
    styles["Normal"]._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), profile.cjk_font)

    # Title
    title_para = doc.add_paragraph(title, style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = profile.latin_font
        run._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), profile.heading_cjk_font)
        run.font.size = Pt(profile.title_pt)

    parser = _TipTapDocxParser()
    parser.feed(html)
    parser.close()
    _render_blocks_to_docx(doc, parser.blocks, quality, profile=profile)
    doc.core_properties.subject = "academic-docx-export"
    doc.core_properties.keywords = quality.to_header()

    doc.save(str(output_path))
    return output_path, quality


@router.post("/docx")
async def export_docx(req: ExportDocxRequest):
    """Export TipTap content as formatted DOCX."""
    action_preflight = _build_docx_action_preflight(req)
    if isinstance(action_preflight, JSONResponse):
        return action_preflight
    tmp_dir = Path(tempfile.mkdtemp(prefix="export_docx_"))
    filename = f"{_safe_docx_filename_stem(req.title)}_{uuid.uuid4().hex[:8]}.docx"
    output_path = tmp_dir / filename

    try:
        _output_path, quality = _html_to_docx(
            req.html,
            req.title,
            output_path,
            req.style_profile,
            req.verify_with_word,
            req.project_id,
        )
    except ValueError as e:
        _cleanup_export_tmp_dir(tmp_dir)
        raise HTTPException(status_code=400, detail=str(e)) from e
    except Exception as e:
        _cleanup_export_tmp_dir(tmp_dir)
        logger.error("Failed to render export DOCX: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to render export DOCX")

    return FileResponse(
        path=str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={
            "X-LitAssist-Export-Quality": quality.to_header(),
            "X-Scholar-AI-Action-Preflight": _preflight_header_value(action_preflight),
            "X-LitAssist-Action-Preflight": _preflight_header_value(action_preflight),
        },
        background=BackgroundTask(_cleanup_export_tmp_dir, tmp_dir),
    )


@router.post("/journal-style-specs/draft")
async def draft_journal_style_spec(req: JournalStyleSpecDraftRequest) -> dict[str, Any]:
    """Create a reviewable journal style profile draft from bounded text."""

    source = {
        "kind": "text",
        "filename": "",
        "bytes": len(req.spec_text.encode("utf-8")),
    }
    try:
        draft = _draft_journal_style_profile(
            project_id=req.project_id,
            journal_name=req.journal_name,
            spec_text=req.spec_text,
            source=source,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    store = _load_project_style_profile_store(req.project_id)
    store["drafts"][draft["draft_id"]] = draft
    _save_project_style_profile_store(req.project_id, store)
    draft["profile"] = _profile_to_public(draft["profile"])
    return draft


@router.post("/journal-style-specs/upload")
async def upload_journal_style_spec(
    project_id: str = Form(..., min_length=1, max_length=128),
    journal_name: str = Form(..., min_length=1, max_length=160),
    file: UploadFile = File(...),
) -> dict[str, Any]:
    """Create a reviewable journal style profile draft from a small text file."""

    filename = Path(str(file.filename or "")).name
    suffix = Path(filename).suffix.lower()
    if suffix not in _JOURNAL_SPEC_ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="Unsupported journal style spec file type")
    media_type = str(file.content_type or "application/octet-stream").lower()
    if media_type not in _JOURNAL_SPEC_ALLOWED_MEDIA_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported journal style spec media type")
    content = await file.read(_JOURNAL_SPEC_MAX_UPLOAD_BYTES + 1)
    if len(content) > _JOURNAL_SPEC_MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Journal style spec file is too large")
    try:
        spec_text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=400, detail="Journal style spec file must be UTF-8 text") from exc
    source = {
        "kind": "upload",
        "filename": filename,
        "bytes": len(content),
    }
    try:
        draft = _draft_journal_style_profile(
            project_id=project_id,
            journal_name=journal_name,
            spec_text=spec_text,
            source=source,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    store = _load_project_style_profile_store(project_id)
    store["drafts"][draft["draft_id"]] = draft
    _save_project_style_profile_store(project_id, store)
    draft["profile"] = _profile_to_public(draft["profile"])
    return draft


@router.post("/journal-style-specs/confirm")
async def confirm_journal_style_spec(req: JournalStyleSpecConfirmRequest) -> dict[str, Any]:
    """Confirm a reviewable style draft for project-scoped DOCX export."""

    store = _load_project_style_profile_store(req.project_id)
    draft = store["drafts"].get(req.draft_id)
    if not isinstance(draft, dict):
        raise HTTPException(status_code=404, detail="Journal style spec draft not found")
    profile = draft.get("profile")
    if not isinstance(profile, dict):
        raise HTTPException(status_code=400, detail="Journal style spec draft is malformed")
    profile_id = str(profile.get("profile_id") or "").strip()
    if not profile_id.startswith("custom_"):
        raise HTTPException(status_code=400, detail="Journal style profile id is invalid")
    confirmed = dict(profile)
    confirmed["confirmed_at"] = _utc_now()
    confirmed["confirmed_by"] = str(req.confirmed_by or "user").strip()[:120]
    store["profiles"][profile_id] = confirmed
    draft["status"] = "confirmed"
    draft["confirmed_at"] = confirmed["confirmed_at"]
    _save_project_style_profile_store(req.project_id, store)
    return {
        "project_id": req.project_id,
        "draft_id": req.draft_id,
        "status": "confirmed",
        "profile": _profile_to_public(confirmed),
        "confirmed_at": confirmed["confirmed_at"],
    }
