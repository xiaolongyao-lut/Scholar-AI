# -*- coding: utf-8 -*-
"""Pure project-export builder helpers."""

from __future__ import annotations

import base64
import json
import re
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Any, Mapping
from urllib.parse import quote

from models import PdfBboxUnit, coerce_pdf_bbox, pdf_bbox_matches_unit

__all__ = [
    "ProjectExportFormat",
    "_strip_citation_tokens",
    "_shorten_export_text",
    "_material_excerpt",
    "_paragraphs_with_offsets",
    "_build_project_academic_export",
    "_build_project_manuscript_markdown",
    "_build_project_markdown_export",
    "_build_project_latex_export",
    "_build_project_docx_export",
    "_build_project_pdf_export",
    "_material_to_csl_json",
    "_citation_tokens_to_pandoc",
    "_build_project_csl_markdown",
    "_build_project_csl_export",
    "_build_project_csl_docx_export",
    "_build_project_csl_latex_export",
    "_build_project_csl_pdf_export",
    "_build_file_export_payload",
    "_markdown_table_cell",
    "_safe_export_filename_stem",
    "_unique_export_file",
]


class ProjectExportFormat(str, __import__("enum").Enum):
    MARKDOWN = "markdown"
    JSON = "json"
    WORD = "word"
    LATEX = "latex"
    PDF = "pdf"


_APP_CITATION_TOKEN_RE = re.compile(r"\[\^cite:([^\]:\]]+)(?::[^\]]*)?\]")


def _strip_citation_tokens(value: str) -> str:
    return re.sub(r"\[\^([^\]]+)\]", "", value).replace("\n", " ").strip()


def _shorten_export_text(value: str, max_length: int = 180) -> str:
    normalized = re.sub(r"\s+", " ", value).strip()
    if len(normalized) <= max_length:
        return normalized
    return f"{normalized[: max_length - 1].rstrip()}…"


def _material_excerpt(material: Any) -> str:
    focus_points = getattr(material, "focus_points", None) or []
    return str(
        getattr(material, "summary", "")
        or (focus_points[0] if focus_points else "")
        or getattr(material, "title", "")
    ).strip()


def _coerce_export_positive_int(value: Any) -> int | None:
    """Return a positive integer metadata value or None."""

    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value if value > 0 else None
    if isinstance(value, str) and value.strip().isdigit():
        parsed = int(value.strip())
        return parsed if parsed > 0 else None
    return None


def _coerce_export_non_negative_int(value: Any) -> int | None:
    """Return a non-negative integer metadata value or None."""

    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value if value >= 0 else None
    if isinstance(value, str) and value.strip().isdigit():
        parsed = int(value.strip())
        return parsed if parsed >= 0 else None
    return None


def _coerce_normalized_export_bbox(value: Any) -> list[float] | None:
    """Return a normalized-ratio bbox suitable for PDF reader URLs."""

    bbox = coerce_pdf_bbox(value)
    if bbox is None:
        return None
    return bbox if pdf_bbox_matches_unit(bbox, PdfBboxUnit.NORMALIZED_RATIO) else None


def _format_export_bbox_param(bbox: list[float]) -> str:
    """Return the compact bbox query-string value used by PDF deep links."""

    return ",".join(str(round(item, 4)).rstrip("0").rstrip(".") for item in bbox)


def _build_source_open_url(
    material_id: str,
    *,
    page: int | None = None,
    chunk_id: str | None = None,
    bbox: list[float] | None = None,
) -> str:
    """Build an app-local URL that opens the original PDF reader."""

    normalized_material_id = str(material_id or "").strip()
    if not normalized_material_id:
        raise ValueError("material_id must be non-empty")
    params: list[tuple[str, str]] = []
    if page is not None:
        params.append(("page", str(page)))
    if chunk_id:
        params.append(("chunk", chunk_id))
    if bbox:
        params.append(("bbox", _format_export_bbox_param(bbox)))
    query = "&".join(f"{quote(key, safe='')}={quote(value, safe='')}" for key, value in params)
    base = f"/workbench/paper/{quote(normalized_material_id, safe='')}"
    return f"{base}?{query}" if query else base


def _build_project_source_anchor(
    material_id: str,
    *,
    chunk_id: Any = None,
    page: Any = None,
    bbox: Any = None,
    text_preview: Any = "",
) -> dict[str, Any] | None:
    """Return a PDF-first source anchor for export appendices."""

    normalized_material_id = str(material_id or "").strip()
    if not normalized_material_id:
        return None
    normalized_chunk_id = str(chunk_id or "").strip() or None
    normalized_page = _coerce_export_positive_int(page)
    normalized_bbox = _coerce_normalized_export_bbox(bbox) if normalized_page is not None else None
    anchor = {
        "material_id": normalized_material_id,
        "chunk_id": normalized_chunk_id,
        "page": normalized_page,
        "bbox": normalized_bbox,
        "bbox_unit": PdfBboxUnit.NORMALIZED_RATIO.value if normalized_bbox is not None else None,
        "text_preview": _shorten_export_text(str(text_preview or ""), 180),
        "open_url": _build_source_open_url(
            normalized_material_id,
            page=normalized_page,
            chunk_id=normalized_chunk_id,
            bbox=normalized_bbox,
        ),
    }
    return anchor


def _chunk_export_preview(chunk: Mapping[str, Any]) -> str:
    """Return a compact source preview from chunk text fields."""

    for key in ("text_preview", "raw_content", "content", "text"):
        value = str(chunk.get(key) or "").strip()
        if value:
            return _shorten_export_text(value, 180)
    return ""


def _sorted_material_chunks(chunk_store: Mapping[str, Any], material_id: str) -> list[dict[str, Any]]:
    """Return material chunks in deterministic chunk-index order."""

    chunks = chunk_store.get(material_id)
    if not isinstance(chunks, list):
        return []
    rows = [dict(chunk) for chunk in chunks if isinstance(chunk, Mapping)]
    return sorted(
        rows,
        key=lambda chunk: (
            _coerce_export_non_negative_int(chunk.get("chunk_index")) or 0,
            str(chunk.get("chunk_id") or ""),
        ),
    )


def _first_material_source_anchor(
    project_id: str | None,
    material_id: str,
    chunk_store: Mapping[str, Any] | None,
) -> dict[str, Any] | None:
    """Resolve the best available export anchor for one material."""

    normalized_material_id = str(material_id or "").strip()
    if not normalized_material_id:
        return None
    if not isinstance(chunk_store, Mapping):
        return _build_project_source_anchor(normalized_material_id)

    for chunk in _sorted_material_chunks(chunk_store, normalized_material_id):
        chunk_id = str(chunk.get("chunk_id") or "").strip()
        if not chunk_id:
            continue
        locator: dict[str, Any] | None = None
        if project_id:
            try:
                from routers.resources_router.endpoints_search_upload import (
                    enrich_chunk_locator_with_pdf,
                    find_chunk_locator,
                )

                located = find_chunk_locator(dict(chunk_store), chunk_id)
                if located is not None:
                    locator = enrich_chunk_locator_with_pdf(str(project_id), dict(chunk_store), located)
            except (ImportError, RuntimeError, TypeError, ValueError):
                locator = None
        if locator is None:
            locator = {
                "material_id": normalized_material_id,
                "chunk_id": chunk_id,
                "page": chunk.get("page"),
                "bbox": chunk.get("bbox"),
            }
        return _build_project_source_anchor(
            str(locator.get("material_id") or normalized_material_id),
            chunk_id=locator.get("chunk_id") or chunk_id,
            page=locator.get("page"),
            bbox=locator.get("bbox"),
            text_preview=locator.get("text_preview") or _chunk_export_preview(chunk),
        )

    return _build_project_source_anchor(normalized_material_id)


def _build_material_source_anchors(
    project_id: str | None,
    materials: list[Any],
    chunk_store: Mapping[str, Any] | None,
) -> dict[str, dict[str, Any]]:
    """Build source anchors keyed by material id without mutating chunk stores."""

    anchors: dict[str, dict[str, Any]] = {}
    for material in materials:
        material_id = str(getattr(material, "material_id", "") or "").strip()
        if not material_id:
            continue
        anchor = _first_material_source_anchor(project_id, material_id, chunk_store)
        if anchor is not None:
            anchors[material_id] = anchor
    return anchors


def _source_anchor_label(anchor: Mapping[str, Any] | None) -> str:
    """Return a compact human-readable source anchor label for appendices."""

    if not isinstance(anchor, Mapping):
        return ""
    parts: list[str] = []
    page = _coerce_export_positive_int(anchor.get("page"))
    if page is not None:
        parts.append(f"p.{page}")
    chunk_id = str(anchor.get("chunk_id") or "").strip()
    if chunk_id:
        parts.append(chunk_id)
    bbox = _coerce_normalized_export_bbox(anchor.get("bbox"))
    if bbox is not None:
        parts.append(f"bbox={_format_export_bbox_param(bbox)}")
    open_url = str(anchor.get("open_url") or "").strip()
    suffix = f" ({'; '.join(parts)})" if parts else ""
    return f"{open_url}{suffix}" if open_url else "; ".join(parts)


def _asset_to_mapping(asset: Any) -> dict[str, Any]:
    """Convert a dataclass-like figure asset to a mapping."""

    if isinstance(asset, Mapping):
        return dict(asset)
    to_dict = getattr(asset, "to_dict", None)
    if callable(to_dict):
        payload = to_dict()
        if isinstance(payload, Mapping):
            return dict(payload)
    return {
        key: getattr(asset, key, None)
        for key in (
            "asset_id",
            "project_id",
            "kind",
            "caption",
            "numbering",
            "material_id",
            "source_page",
            "bbox",
            "asset_path",
            "width",
            "height",
            "format",
        )
    }


def _build_project_figure_assets_export(
    project_id: str | None,
    figure_assets: list[Any] | None,
) -> list[dict[str, Any]]:
    """Return export-safe figure/table asset provenance rows."""

    if not figure_assets:
        return []
    rows: list[dict[str, Any]] = []
    for asset in figure_assets:
        payload = _asset_to_mapping(asset)
        asset_id = str(payload.get("asset_id") or "").strip()
        asset_path = str(payload.get("asset_path") or "").strip()
        kind = str(payload.get("kind") or "").strip()
        caption = str(payload.get("caption") or "").strip()
        numbering = str(payload.get("numbering") or "").strip()
        if not asset_id or kind not in {"figure", "table"} or not caption or not numbering or not asset_path:
            continue
        material_id = str(payload.get("material_id") or "").strip() or None
        page = _coerce_export_positive_int(payload.get("source_page"))
        bbox = coerce_pdf_bbox(payload.get("bbox"))
        normalized_bbox = _coerce_normalized_export_bbox(bbox) if page is not None else None
        source_anchor = (
            _build_project_source_anchor(
                material_id,
                page=page,
                bbox=normalized_bbox,
                text_preview=caption,
            )
            if material_id
            else None
        )
        rows.append(
            {
                "asset_id": asset_id,
                "project_id": str(payload.get("project_id") or project_id or ""),
                "kind": kind,
                "caption": caption,
                "numbering": numbering,
                "material_id": material_id,
                "source_page": page,
                "bbox": bbox,
                "bbox_unit": PdfBboxUnit.NORMALIZED_RATIO.value if normalized_bbox is not None else None,
                "asset_path": asset_path,
                "width": _coerce_export_positive_int(payload.get("width")),
                "height": _coerce_export_positive_int(payload.get("height")),
                "format": str(payload.get("format") or "").strip() or None,
                "source_anchor": source_anchor,
            }
        )
    return sorted(rows, key=lambda row: (str(row["kind"]), str(row["numbering"]), str(row["asset_id"])))


def _material_metadata_value(material: Any, key: str) -> Any:
    metadata = getattr(material, "metadata", None)
    if isinstance(metadata, Mapping):
        return metadata.get(key)
    return None


def _string_list_metadata(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        items = re.split(r"\s*(?:;|,|、)\s*", value)
    elif isinstance(value, (list, tuple)):
        items = [str(item) for item in value]
    else:
        return []
    return [item.strip() for item in items if item.strip()]


def _build_bibliography_entries(materials: list[Any]) -> list[dict[str, Any]]:
    """Build deterministic reference entries from available material metadata."""
    entries: list[dict[str, Any]] = []
    for ordinal, material in enumerate(materials, 1):
        title = str(getattr(material, "title", "") or getattr(material, "title_en", "") or "无标题").strip()
        authors = _string_list_metadata(_material_metadata_value(material, "authors"))
        year = str(_material_metadata_value(material, "year") or "").strip() or None
        venue = str(_material_metadata_value(material, "venue") or "").strip() or None
        doi = str(_material_metadata_value(material, "doi") or "").strip() or None
        url = str(_material_metadata_value(material, "url") or "").strip() or None
        prefix = ", ".join(authors) if authors else ""
        year_part = f" ({year})" if year and prefix else ""
        venue_part = f". {venue}" if venue else ""
        doi_part = f". DOI: {doi}" if doi else ""
        url_part = f". {url}" if url and not doi else ""
        title_part = f". {title}" if prefix else title
        display_text = f"{prefix}{year_part}{title_part}{venue_part}{doi_part}{url_part}."
        display_text = re.sub(r"\s+", " ", display_text).replace("..", ".").strip()
        entries.append(
            {
                "citation_key": f"ref-{ordinal}",
                "material_id": material.material_id,
                "ordinal": ordinal,
                "title": title,
                "type": str(getattr(material, "type", "reference") or "reference"),
                "authors": authors,
                "year": year,
                "venue": venue,
                "doi": doi,
                "url": url,
                "summary": str(getattr(material, "summary", "") or ""),
                "display_text": display_text,
            }
        )
    return entries


def _append_unique_material_id(material_ids: list[str], seen: set[str], value: Any) -> None:
    """Append one non-empty material id while preserving first-seen order."""

    material_id = str(value or "").strip()
    if not material_id or material_id in seen:
        return
    seen.add(material_id)
    material_ids.append(material_id)


def _cited_material_ids_from_drafts(drafts: list[Any]) -> list[str]:
    """Return material ids that are actually cited in draft content or anchors."""

    material_ids: list[str] = []
    seen: set[str] = set()
    for draft in drafts:
        content = str(getattr(draft, "content", "") or "")
        for match in _APP_CITATION_TOKEN_RE.finditer(content):
            _append_unique_material_id(material_ids, seen, match.group(1))
        to_dict = getattr(draft, "to_dict", None)
        payload = to_dict() if callable(to_dict) else {}
        anchors = payload.get("citation_anchors") if isinstance(payload, Mapping) else []
        if not isinstance(anchors, list):
            continue
        for anchor in anchors:
            if not isinstance(anchor, Mapping):
                continue
            _append_unique_material_id(material_ids, seen, anchor.get("materialId"))
    return material_ids


def _cited_materials(materials: list[Any], drafts: list[Any]) -> list[Any]:
    """Return material records cited by manuscript drafts in citation order."""

    material_by_id = {str(getattr(material, "material_id", "") or ""): material for material in materials}
    return [
        material_by_id[material_id]
        for material_id in _cited_material_ids_from_drafts(drafts)
        if material_id in material_by_id
    ]


def _citation_number_map(drafts: list[Any]) -> dict[str, int]:
    """Return stable one-based citation numbers by cited material id."""

    return {
        material_id: index
        for index, material_id in enumerate(_cited_material_ids_from_drafts(drafts), 1)
    }


def _render_manuscript_citations_as_numbers(content: str, citation_numbers: Mapping[str, int]) -> str:
    """Render app citation tokens as compact numeric manuscript citations."""

    if not isinstance(content, str):
        raise TypeError("content must be a string")

    def replace(match: re.Match[str]) -> str:
        material_id = match.group(1)
        number = citation_numbers.get(material_id)
        return f"[{number}]" if number is not None else ""

    return _APP_CITATION_TOKEN_RE.sub(replace, content).strip()


def _paragraphs_with_offsets(content: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    separator = re.compile(r"\n\s*\n+")
    last_index = 0

    def push(raw_segment: str, raw_start: int, raw_end: int) -> None:
        if not raw_segment.strip():
            return
        leading = len(raw_segment) - len(raw_segment.lstrip())
        trailing = len(raw_segment) - len(raw_segment.rstrip())
        start_offset = raw_start + leading
        end_offset = max(start_offset, raw_end - trailing)
        records.append(
            {
                "index": len(records) + 1,
                "text": raw_segment.strip(),
                "start_offset": start_offset,
                "end_offset": end_offset,
            }
        )

    for match in separator.finditer(content):
        push(content[last_index:match.start()], last_index, match.start())
        last_index = match.end()
    push(content[last_index:], last_index, len(content))
    return records


def _build_project_academic_export(
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    *,
    project_id: str | None = None,
    chunk_store: Mapping[str, Any] | None = None,
    figure_assets: list[Any] | None = None,
) -> dict[str, list[dict[str, Any]]]:
    """Derive academic evidence view-models from existing materials and anchors."""
    material_lookup = {material.material_id: material for material in materials}
    source_anchors = _build_material_source_anchors(project_id, materials, chunk_store)
    anchors_by_material: dict[str, list[dict[str, Any]]] = {}
    citation_chain: list[dict[str, Any]] = []
    review_findings: list[dict[str, Any]] = []
    section_ids = {section.section_id for section in sections}

    for draft in drafts:
        draft_payload = draft.to_dict()
        anchors = draft_payload.get("citation_anchors", [])
        paragraphs = _paragraphs_with_offsets(str(getattr(draft, "content", "")))
        for anchor in anchors:
            if not isinstance(anchor, Mapping):
                continue
            material_id = anchor.get("materialId")
            anchor_id = str(anchor.get("id", "")).strip()
            if material_id:
                anchors_by_material.setdefault(str(material_id), []).append(anchor)
            paragraph = next(
                (
                    item
                    for item in paragraphs
                    if int(anchor.get("startOffset", -1)) >= item["start_offset"]
                    and int(anchor.get("endOffset", -1)) <= item["end_offset"]
                ),
                None,
            )
            material = material_lookup.get(str(material_id)) if material_id else None
            excerpt = _material_excerpt(material) if material else ""
            source_anchor = (
                source_anchors.get(material.material_id)
                if material
                else _build_project_source_anchor(str(material_id), text_preview=excerpt)
                if material_id
                else None
            )
            citation_chain.append(
                {
                    "anchor_id": anchor_id,
                    "section_id": draft.section_id if draft.section_id in section_ids else None,
                    "paragraph_index": paragraph["index"] if paragraph else None,
                    "material_id": material.material_id if material else material_id,
                    "evidence_id": f"evidence:{material.material_id}" if material else None,
                    "claim_excerpt": (
                        _shorten_export_text(_strip_citation_tokens(paragraph["text"]))
                        if paragraph
                        else ""
                    ),
                    "source_excerpt": _shorten_export_text(excerpt) if excerpt else "",
                    "page": source_anchor.get("page") if source_anchor else None,
                    "source_anchor": source_anchor,
                    "confidence": None,
                }
            )

        uncited_long = [
            paragraph
            for paragraph in paragraphs
            if len(_strip_citation_tokens(paragraph["text"])) >= 80
            and not any(
                int(anchor.get("startOffset", -1)) >= paragraph["start_offset"]
                and int(anchor.get("endOffset", -1)) <= paragraph["end_offset"]
                for anchor in anchors
                if isinstance(anchor, Mapping)
            )
        ]
        if uncited_long:
            review_findings.append(
                {
                    "id": f"uncited-paragraphs:{draft.draft_id}",
                    "severity": "warning",
                    "message": f"{len(uncited_long)} long paragraph(s) have no citation anchors.",
                    "draft_id": draft.draft_id,
                    "section_id": draft.section_id,
                }
            )

    for material_id in sorted(anchors_by_material):
        if material_id not in material_lookup:
            review_findings.append(
                {
                    "id": f"dangling-material:{material_id}",
                    "severity": "warning",
                    "message": "Citation anchor points to a material that is not in this project export.",
                    "material_id": material_id,
                }
            )

    evidence_rows: list[dict[str, Any]] = []
    for material in materials:
        excerpt = _material_excerpt(material)
        source_anchor = source_anchors.get(material.material_id) or _build_project_source_anchor(
            material.material_id,
            text_preview=excerpt,
        )
        anchor_ids = [
            str(anchor.get("id", ""))
            for anchor in anchors_by_material.get(material.material_id, [])
        ]
        status = "unused"
        if anchor_ids:
            status = "used" if excerpt else "weak"
        evidence_rows.append(
            {
                "evidence_id": f"evidence:{material.material_id}",
                "material_id": material.material_id,
                "chunk_id": source_anchor.get("chunk_id") if source_anchor else None,
                "page": source_anchor.get("page") if source_anchor else None,
                "excerpt": _shorten_export_text(excerpt),
                "score": None,
                "provenance": {
                    "material_title": material.title,
                    "material_type": material.type,
                },
                "anchor_ids": anchor_ids,
                "source_anchor": source_anchor,
                "status": status,
            }
        )

    bibliography_entries = _build_bibliography_entries(materials)
    figure_asset_rows = _build_project_figure_assets_export(project_id, figure_assets)

    return {
        "evidence_rows": evidence_rows,
        "citation_chain": citation_chain,
        "bibliography_entries": bibliography_entries,
        "review_findings": review_findings,
        "figure_assets": figure_asset_rows,
    }


def _markdown_table_cell(value: Any) -> str:
    return str(value or "").replace("|", "\\|").replace("\n", " ").strip()


def _safe_export_filename_stem(value: str, fallback: str = "writing-project") -> str:
    """Return a filesystem-safe export stem with user text preserved when valid."""
    normalized = re.sub(r"\s+", " ", str(value or "").strip())
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", normalized).strip(" ._")
    if not safe:
        safe = fallback
    return safe[:96]


def _iter_project_export_markdown_lines(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    academic_export: Mapping[str, list[dict[str, Any]]],
) -> list[str]:
    """Build canonical markdown lines used by all writing export formats."""
    lines = [f"# {project.title}\n"]
    if project.description:
        lines.append(f"> {project.description}\n")
    lines.append(f"状态: {project.status} | 创建: {project.created_at}\n")

    sorted_sections = sorted(sections, key=lambda s: s.order)
    section_map = {s.section_id: s for s in sorted_sections}
    material_map = {m.material_id: m for m in materials}

    for section in sorted_sections:
        lines.append(f"\n## {section.title}\n")
        if section.description:
            lines.append(f"{section.description}\n")
        section_drafts = [d for d in drafts if getattr(d, "section_id", None) == section.section_id]
        for draft in section_drafts:
            lines.append(f"\n### {draft.title}\n")
            lines.append(f"{draft.content}\n")

    orphans = [d for d in drafts if not getattr(d, "section_id", None)]
    if orphans:
        lines.append("\n## 未分类草稿\n")
        for draft in orphans:
            lines.append(f"\n### {draft.title}\n")
            lines.append(f"{draft.content}\n")

    evidence_rows = academic_export.get("evidence_rows", [])
    if evidence_rows:
        lines.append("\n## 证据表\n")
        lines.append("| Evidence ID | Material | Status | Anchors | PDF Anchor | Excerpt |")
        lines.append("|---|---|---|---|---|---|")
        for row in evidence_rows:
            anchors = ", ".join(row["anchor_ids"])
            material_title = row["provenance"]["material_title"]
            lines.append(
                "| "
                + " | ".join(
                    [
                        _markdown_table_cell(row["evidence_id"]),
                        _markdown_table_cell(material_title),
                        _markdown_table_cell(row["status"]),
                        _markdown_table_cell(anchors),
                        _markdown_table_cell(_source_anchor_label(row.get("source_anchor"))),
                        _markdown_table_cell(row["excerpt"]),
                    ]
                )
                + " |"
            )

    citation_chain = academic_export.get("citation_chain", [])
    if citation_chain:
        lines.append("\n## 引用链\n")
        lines.append("| Anchor | Section | Paragraph | Material | PDF Anchor | Claim | Source |")
        lines.append("|---|---|---|---|---|---|---|")
        for row in citation_chain:
            section = section_map.get(row["section_id"])
            material = material_map.get(row["material_id"])
            lines.append(
                "| "
                + " | ".join(
                    [
                        _markdown_table_cell(row["anchor_id"]),
                        _markdown_table_cell(section.title if section else ""),
                        _markdown_table_cell(row["paragraph_index"]),
                        _markdown_table_cell(
                            material.title if material else row["material_id"]
                        ),
                        _markdown_table_cell(_source_anchor_label(row.get("source_anchor"))),
                        _markdown_table_cell(row["claim_excerpt"]),
                        _markdown_table_cell(row["source_excerpt"]),
                    ]
                )
                + " |"
            )

    figure_assets = academic_export.get("figure_assets", [])
    if figure_assets:
        lines.append("\n## 图表资产 provenance\n")
        lines.append("| Asset | Kind | Numbering | Source | Page | Bbox | Asset Path |")
        lines.append("|---|---|---|---|---|---|---|")
        for row in figure_assets:
            bbox = row.get("bbox")
            bbox_label = _format_export_bbox_param(bbox) if isinstance(bbox, list) else ""
            lines.append(
                "| "
                + " | ".join(
                    [
                        _markdown_table_cell(row.get("asset_id")),
                        _markdown_table_cell(row.get("kind")),
                        _markdown_table_cell(row.get("numbering")),
                        _markdown_table_cell(_source_anchor_label(row.get("source_anchor"))),
                        _markdown_table_cell(row.get("source_page")),
                        _markdown_table_cell(bbox_label),
                        _markdown_table_cell(row.get("asset_path")),
                    ]
                )
                + " |"
            )

    review_findings = academic_export.get("review_findings", [])
    if review_findings:
        lines.append("\n## 审计提示\n")
        for finding in review_findings:
            lines.append(f"- {finding['message']}")

    bibliography_entries = academic_export.get("bibliography_entries", [])
    if bibliography_entries:
        lines.append("\n## 参考文献\n")
        for entry in bibliography_entries:
            lines.append(f"{entry['ordinal']}. {entry['display_text']}")
            if entry.get("summary"):
                lines.append(f"   摘要: {str(entry['summary'])[:100]}...")
            lines.append("")

    return lines


def _render_manuscript_draft_content(
    draft: Any,
    *,
    citation_style: str,
    citation_numbers: Mapping[str, int],
) -> str:
    """Render one draft as manuscript text without internal app metadata."""

    content = str(getattr(draft, "content", "") or "")
    if citation_style == "pandoc":
        return _citation_tokens_to_pandoc(content).strip()
    if citation_style == "numeric":
        return _render_manuscript_citations_as_numbers(content, citation_numbers)
    raise ValueError("citation_style must be 'numeric' or 'pandoc'")


def _iter_project_manuscript_markdown_lines(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    *,
    citation_style: str = "numeric",
    include_bibliography: bool = True,
    include_title: bool = True,
) -> list[str]:
    """Build manuscript-only Markdown lines for preview and document export."""

    if project is None:
        raise ValueError("project is required")
    if citation_style not in {"numeric", "pandoc"}:
        raise ValueError("citation_style must be 'numeric' or 'pandoc'")
    citation_numbers = _citation_number_map(drafts)
    sorted_sections = sorted(sections, key=lambda section: getattr(section, "order", 0))
    lines: list[str] = []
    title = str(getattr(project, "title", "") or "").strip()
    if include_title and title:
        lines.extend([f"# {title}", ""])

    for section in sorted_sections:
        section_drafts = [
            draft
            for draft in drafts
            if getattr(draft, "section_id", None) == getattr(section, "section_id", None)
            and _render_manuscript_draft_content(
                draft,
                citation_style=citation_style,
                citation_numbers=citation_numbers,
            ).strip()
        ]
        section_title = str(getattr(section, "title", "") or "").strip()
        if section_title:
            lines.extend([f"## {section_title}", ""])
        if not section_drafts:
            continue
        for draft in section_drafts:
            draft_text = _render_manuscript_draft_content(
                draft,
                citation_style=citation_style,
                citation_numbers=citation_numbers,
            )
            draft_title = str(getattr(draft, "title", "") or "").strip()
            if len(section_drafts) > 1 and draft_title:
                lines.extend([f"### {draft_title}", ""])
            lines.extend([draft_text, ""])

    orphans = [
        draft
        for draft in drafts
        if not getattr(draft, "section_id", None)
        and _render_manuscript_draft_content(
            draft,
            citation_style=citation_style,
            citation_numbers=citation_numbers,
        ).strip()
    ]
    for draft in orphans:
        draft_title = str(getattr(draft, "title", "") or "").strip()
        if draft_title:
            lines.extend([f"## {draft_title}", ""])
        lines.extend(
            [
                _render_manuscript_draft_content(
                    draft,
                    citation_style=citation_style,
                    citation_numbers=citation_numbers,
                ),
                "",
            ]
        )

    cited_materials = _cited_materials(materials, drafts)
    if include_bibliography and cited_materials:
        if citation_style == "pandoc":
            lines.extend(["## 参考文献", "", "::: {#refs}", ":::", ""])
        else:
            lines.extend(["## 参考文献", ""])
            for entry in _build_bibliography_entries(cited_materials):
                lines.extend([f"{entry['ordinal']}. {entry['display_text']}", ""])

    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _build_project_manuscript_markdown(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
) -> str:
    """Render project drafts as manuscript-only Markdown without audit appendices."""

    if project is None:
        raise ValueError("project is required")
    return "\n".join(
        _iter_project_manuscript_markdown_lines(project, sections, drafts, materials)
    )


def _build_project_markdown_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    academic_export: Mapping[str, list[dict[str, Any]]],
) -> str:
    """Render the public API Markdown export without internal audit appendices."""

    if project is None:
        raise ValueError("project is required")
    return _build_project_manuscript_markdown(project, sections, drafts, materials)


def _latex_escape(value: Any) -> str:
    """Escape plain text for a LaTeX document body."""
    text = str(value or "")
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in text)


def _build_project_latex_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    academic_export: Mapping[str, list[dict[str, Any]]],
) -> str:
    """Render a standalone LaTeX article from manuscript-only project drafts."""
    manuscript_lines = _iter_project_manuscript_markdown_lines(project, sections, drafts, materials)
    title = str(getattr(project, "title", "") or "").strip()
    lines = [
        r"\documentclass[UTF8]{ctexart}",
        r"\usepackage{geometry}",
        r"\geometry{a4paper, margin=2.5cm}",
    ]
    if title:
        lines.extend(
            [
                f"\\title{{{_latex_escape(title)}}}",
                r"\date{}",
            ]
        )
    lines.extend(
        [
        r"\begin{document}",
        "",
        ]
    )
    if title:
        lines.extend([r"\maketitle", ""])
    for raw_line in manuscript_lines:
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("### "):
            lines.extend([f"\\subsection{{{_latex_escape(stripped[4:])}}}", ""])
            continue
        if stripped.startswith("## "):
            lines.extend([f"\\section{{{_latex_escape(stripped[3:])}}}", ""])
            continue
        for paragraph in _paragraphs_with_offsets(stripped):
            lines.extend([_latex_escape(paragraph["text"]), ""])

    lines.append(r"\end{document}")
    return "\n".join(lines)


def _add_docx_paragraphs_from_text(document: Any, text: str) -> None:
    """Append non-empty paragraphs to a python-docx document."""
    for paragraph in _paragraphs_with_offsets(text):
        document.add_paragraph(paragraph["text"])


def _build_project_docx_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    academic_export: Mapping[str, list[dict[str, Any]]],
    output_path: Path,
) -> Path:
    """Write a DOCX manuscript export and return its path."""
    if output_path.suffix.lower() != ".docx":
        raise ValueError("output_path must use a .docx suffix")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word export") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    for raw_line in _iter_project_manuscript_markdown_lines(project, sections, drafts, materials):
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            document.add_paragraph(number_match.group(1).strip(), style="List Number")
            continue
        _add_docx_paragraphs_from_text(document, stripped)

    document.save(str(output_path))
    return output_path


def _split_csl_author(name: str) -> dict[str, str]:
    """Parse an author string into CSL name parts.

    "Family, Given" splits; otherwise use ``literal`` so Chinese names and
    pre-formatted Western names render verbatim instead of being mis-split.
    """
    trimmed = name.strip()
    comma = trimmed.find(",")
    if comma > 0:
        return {"family": trimmed[:comma].strip(), "given": trimmed[comma + 1:].strip()}
    return {"literal": trimmed}


def _material_to_csl_json(material: Any) -> dict[str, Any]:
    """Map a material + its persisted bibliographic metadata to a CSL-JSON item.

    The material_id is the stable CSL id (and pandoc citation key). Fields come
    from ``material.metadata`` under the same keys the citation manager writes.
    """
    metadata = getattr(material, "metadata", None)
    meta: Mapping[str, Any] = metadata if isinstance(metadata, Mapping) else {}
    item: dict[str, Any] = {
        "id": str(getattr(material, "material_id", "") or ""),
        "type": str(meta.get("csl_type") or "article-journal").strip() or "article-journal",
        "title": str(getattr(material, "title", "") or "").strip(),
        "language": "zh-CN",
    }
    authors = [_split_csl_author(a) for a in _string_list_metadata(meta.get("authors")) if a.strip()]
    if authors:
        item["author"] = authors
    year_raw = str(meta.get("year") or "").strip()
    if year_raw:
        try:
            item["issued"] = {"date-parts": [[int(year_raw)]]}
        except ValueError:
            pass
    for csl_key, meta_key in (
        ("container-title", "venue"),
        ("publisher", "publisher"),
        ("volume", "volume"),
        ("issue", "issue"),
        ("page", "pages"),
        ("DOI", "doi"),
        ("URL", "url"),
    ):
        value = str(meta.get(meta_key) or "").strip()
        if value:
            item[csl_key] = value
    return item


def _citation_tokens_to_pandoc(text: str) -> str:
    """Convert ``[^cite:material_id:anchor]`` draft tokens to pandoc ``[@material_id]``."""
    return re.sub(r"\[\^cite:([^\]:]+)(?::[^\]]*)?\]", r"[@\1]", str(text or ""))


def _build_project_csl_markdown(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
) -> str:
    """Build pandoc Markdown from manuscript drafts with CSL cite tokens only."""
    if project is None:
        raise ValueError("project is required")
    title = str(getattr(project, "title", "") or "").strip()
    body = "\n".join(
        _iter_project_manuscript_markdown_lines(
            project,
            sections,
            drafts,
            materials,
            citation_style="pandoc",
            include_bibliography=True,
            include_title=False,
        )
    )
    if not title:
        return body
    return "\n".join(["---", f"title: {json.dumps(title, ensure_ascii=False)}", "---", "", body]).strip()


def _build_project_csl_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    style_xml: str,
    output_path: Path,
    *,
    extra_args: list[str] | None = None,
) -> Path:
    """Render a project export with CSL-formatted in-text citations + bibliography via pandoc.

    pandoc infers the target format from ``output_path``'s suffix (``.docx`` /
    ``.tex`` / ``.pdf`` …). Uses the active CSL style (``style_xml``) and a
    CSL-JSON bibliography derived from material metadata. Raises ``RuntimeError``
    when pandoc is unavailable, no style is provided, or the conversion fails, so
    callers can fall back to a deterministic builder.
    """
    if not str(style_xml or "").strip():
        raise RuntimeError("无可用 CSL 样式")
    if not shutil.which("pandoc"):
        raise RuntimeError("pandoc 未安装")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    items = [_material_to_csl_json(material) for material in _cited_materials(materials, drafts)]
    body = _build_project_csl_markdown(project, sections, drafts, materials)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        (tmp_dir / "style.csl").write_text(style_xml, encoding="utf-8")
        (tmp_dir / "refs.json").write_text(json.dumps(items, ensure_ascii=False), encoding="utf-8")
        (tmp_dir / "input.md").write_text(body, encoding="utf-8")
        command = [
            "pandoc", "input.md",
            "--citeproc",
            "--csl", "style.csl",
            "--bibliography", "refs.json",
            "--metadata", "lang=zh-CN",
        ]
        if extra_args:
            command += list(extra_args)
        command += ["-o", str(output_path)]
        try:
            result = subprocess.run(
                command,
                cwd=str(tmp_dir),
                capture_output=True,
                text=True,
                timeout=180,
            )
        except (OSError, subprocess.SubprocessError) as exc:
            raise RuntimeError(f"pandoc 调用失败: {exc}") from exc
        if result.returncode != 0:
            raise RuntimeError(f"pandoc 转换失败: {(result.stderr or '').strip()[:400]}")
    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("pandoc 未生成有效输出文件")
    return output_path


def _build_project_csl_docx_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    style_xml: str,
    output_path: Path,
) -> Path:
    """Render a DOCX with CSL-formatted citations + bibliography (see _build_project_csl_export)."""
    if output_path.suffix.lower() != ".docx":
        raise ValueError("output_path must use a .docx suffix")
    return _build_project_csl_export(project, sections, drafts, materials, style_xml, output_path)


def _build_project_csl_latex_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    style_xml: str,
    output_path: Path,
) -> Path:
    """Render a standalone LaTeX (.tex) document with CSL-formatted citations + bibliography.

    pandoc renders ``.tex`` text without needing a TeX engine, so this is safe
    wherever pandoc is installed.
    """
    if output_path.suffix.lower() != ".tex":
        raise ValueError("output_path must use a .tex suffix")
    return _build_project_csl_export(
        project, sections, drafts, materials, style_xml, output_path, extra_args=["--standalone"]
    )


def _build_project_csl_pdf_export(
    project: Any,
    sections: list[Any],
    drafts: list[Any],
    materials: list[Any],
    style_xml: str,
    output_path: Path,
) -> Path:
    """Render a PDF with CSL-formatted citations + bibliography via pandoc.

    Requires a PDF engine (e.g. xelatex for CJK content); raises ``RuntimeError``
    when the engine is missing so the caller can fall back to the PyMuPDF text
    PDF builder.
    """
    if output_path.suffix.lower() != ".pdf":
        raise ValueError("output_path must use a .pdf suffix")
    return _build_project_csl_export(project, sections, drafts, materials, style_xml, output_path)


def _build_project_pdf_export(markdown_content: str, output_path: Path, title: str) -> Path:
    """Write a simple text PDF export using PyMuPDF."""
    if output_path.suffix.lower() != ".pdf":
        raise ValueError("output_path must use a .pdf suffix")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required for PDF export") from exc

    document = fitz.open()
    page_width = 595
    page_height = 842
    margin = 54
    font_size = 10.5
    line_height = 15
    max_chars = 78
    y = margin
    page = document.new_page(width=page_width, height=page_height)

    def add_line(line: str, *, size: float = font_size) -> None:
        nonlocal page, y
        if y > page_height - margin:
            page = document.new_page(width=page_width, height=page_height)
            y = margin
        page.insert_text((margin, y), line, fontsize=size, fontname="helv")
        y += line_height if size <= font_size else line_height + 4

    add_line(str(title), size=16)
    add_line("")
    for raw_line in markdown_content.splitlines():
        line = raw_line.strip()
        if not line:
            add_line("")
            continue
        while len(line) > max_chars:
            add_line(line[:max_chars])
            line = line[max_chars:]
        add_line(line)

    document.save(str(output_path))
    document.close()
    return output_path


def _build_file_export_payload(
    *,
    project_id: str,
    format_name: str,
    filename: str,
    file_path: Path,
    media_type: str,
    content: str | None = None,
    evidence_rows: list[dict[str, Any]] | None = None,
    citation_chain: list[dict[str, Any]] | None = None,
    bibliography_entries: list[dict[str, Any]] | None = None,
    review_findings: list[dict[str, Any]] | None = None,
    figure_assets: list[dict[str, Any]] | None = None,
    writing_audit: Mapping[str, Any] | None = None,
    rendered_writing_audit: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Return a JSON-safe file payload containing path metadata and base64."""
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(str(file_path))
    encoded = base64.b64encode(file_path.read_bytes()).decode("ascii")
    return {
        "project_id": project_id,
        "format": format_name,
        "filename": filename,
        "content": content,
        "content_base64": encoded,
        "media_type": media_type,
        "file_path": str(file_path),
        "evidence_rows": evidence_rows or [],
        "citation_chain": citation_chain or [],
        "bibliography_entries": bibliography_entries or [],
        "review_findings": review_findings or [],
        "figure_assets": figure_assets or [],
        "writing_audit": dict(writing_audit) if isinstance(writing_audit, Mapping) else None,
        "rendered_writing_audit": (
            dict(rendered_writing_audit) if isinstance(rendered_writing_audit, Mapping) else None
        ),
    }


def _unique_export_file(output_dir: Path, stem: str, suffix: str) -> Path:
    """Return a collision-resistant generated export path."""
    return output_dir / f"{stem}-{uuid.uuid4().hex[:8]}{suffix}"
