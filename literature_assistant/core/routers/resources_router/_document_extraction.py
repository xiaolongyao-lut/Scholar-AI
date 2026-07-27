# -*- coding: utf-8 -*-
"""Pure document-content extraction helpers."""

from __future__ import annotations

import json
import logging
import hashlib
import io
import math
import re
import unicodedata
from collections import Counter, OrderedDict
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from pathlib import Path
from threading import Lock
from typing import TYPE_CHECKING, Any, cast

from ._pymupdf_dynamic import (
    new_pymupdf_matrix,
    new_pymupdf_rect,
    open_pymupdf_document,
    pymupdf_rect_area,
    transform_pymupdf_rect,
)

if TYPE_CHECKING:
    from literature_assistant.core.pdf_backends import (
        PDFParserProvenance,
        StructuredBlock,
        get_pdf_backend,
        parse_pdf_with_provenance,
    )
    from literature_assistant.core.pdf_backends.ocr_ingestion import apply_pdf_ocr_if_needed
    from literature_assistant.core.pdf_backends.pymupdf_backend import PyMuPDFBackend
else:
    try:
        from pdf_backends import (
            PDFParserProvenance,
            StructuredBlock,
            get_pdf_backend,
            parse_pdf_with_provenance,
        )
        from pdf_backends.ocr_ingestion import apply_pdf_ocr_if_needed
        from pdf_backends.pymupdf_backend import PyMuPDFBackend
    except ImportError:  # pragma: no cover — only triggered in misconfigured envs
        StructuredBlock = None  # type: ignore[assignment]
        PDFParserProvenance = None  # type: ignore[assignment]
        get_pdf_backend = None  # type: ignore[assignment]
        parse_pdf_with_provenance = None  # type: ignore[assignment]
        apply_pdf_ocr_if_needed = None  # type: ignore[assignment]
        PyMuPDFBackend = None  # type: ignore[assignment]


__all__ = [
    "_extract_document_content",
    "_extract_document_content_from_path",
    "_extract_document_payload_from_path",
    "_reconcile_document_blocks",
    "_truncate_document_content",
    "ExtractedDocumentPayload",
    "PdfFormulaCandidate",
    "bind_pdf_formula_candidates_to_chunks",
    "extract_pymupdf_formula_candidates",
    "formula_candidates_from_chunks",
    "merge_pdf_formula_candidates",
]


_LOGGER = logging.getLogger("DocumentExtraction")
_OCR_CONTENT_MARKER_RE = re.compile(r"^\[OCR Page \d+\]$", re.IGNORECASE)
_PDF_VISUAL_MIN_WIDTH = 96.0
_PDF_VISUAL_MIN_HEIGHT = 72.0
_PDF_VISUAL_MIN_AREA_RATIO = 0.01
_BROWSER_SAFE_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
_IMAGE_EXTENSION_ALIASES = {"jpe": "jpg", "jfif": "jpg"}
_PDF_FIGURE_TABLE_REF_RE = re.compile(
    r"(?P<prefix>图|圖|表|fig(?:ure)?\.?|table)\s*"
    r"(?P<number>[A-Za-z]?\d+(?:[.\-–—]\d+)*[A-Za-z]?)",
    re.IGNORECASE,
)
_PDF_BODY_REF_VERB_RE = re.compile(
    r"^(?:shows?|presents?|illustrates?|demonstrates?|reports?|describes?|"
    r"compares?|summari[sz]es?|lists?|indicates?|reveals?|suggests?|"
    r"can|is|are|was|were|has|have|had)\b",
    re.IGNORECASE,
)
_PDF_BODY_REF_CJK_VERB_RE = re.compile(
    r"^(?:中|所示|所列|显示|表明|说明|展示|呈现|给出|列出|总结|比较|报道|描述|揭示|指出|可见|可知|可以|能够|为|是|有)"
)
_PDF_FORMULA_SCAN_LIMIT = 200
_PDF_FORMULA_CACHE_MAX = 24
_PDF_FORMULA_FRAGMENT_MAX_PER_PAGE = 256
_PDF_FORMULA_ANCHOR_MAX_PER_PAGE = 64
_PDF_FORMULA_DETECTOR_VERSION = "fragment-v3"
_PDF_FORMULA_TEXT_MAX_CHARS = 512
_PDF_FORMULA_NUMBER_TOKEN = r"\d{1,3}(?:[.\-–—]\d{1,3})?[A-Za-z]?"
_PDF_FORMULA_NUMBER_ONLY_RE = re.compile(
    rf"^[（(\[]\s*{_PDF_FORMULA_NUMBER_TOKEN}\s*[）)\]]\s*[,.;:]?$"
)
_PDF_FORMULA_TRAILING_NUMBER_RE = re.compile(
    rf"\s*[（(\[]\s*{_PDF_FORMULA_NUMBER_TOKEN}\s*[）)\]]\s*[,.;:]?$"
)
_PDF_FORMULA_RELATION_RE = re.compile(r"[=≠≈≃≅≡≤≥<>∝]")
_PDF_FORMULA_STRONG_MATH_RE = re.compile(
    r"[∑∏∫∮√∞∂∇⊗⊕→←↔⇒⇔∈∉⊂⊆⊃⊇∪∩∀∃]"
)
_PDF_FORMULA_OPERATOR_RE = re.compile(r"[+\-−*/^±∓×÷·]")
_PDF_FORMULA_GREEK_RE = re.compile(r"[\u0370-\u03ff\u1f00-\u1fff]")
_PDF_FORMULA_SCRIPT_RE = re.compile(
    r"[⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ]"
)
_PDF_FORMULA_URL_RE = re.compile(
    r"(?:https?://|www\.|\bdoi\s*:|\b10\.\d{4,9}/\S+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,})",
    re.IGNORECASE,
)
_PDF_FORMULA_CAPTION_RE = re.compile(
    r"^(?:fig(?:ure)?\.?|table|图|圖|表)\s*[A-Za-z]?\d+",
    re.IGNORECASE,
)
_PDF_FORMULA_LIST_RE = re.compile(r"^(?:[-*•]|\d+[.)、])\s+")
_PDF_FORMULA_DISCOURSE_RE = re.compile(
    r"^(?:where|when|with|let|thus|hence|therefore|其中|式中|当|若|则|令)\b",
    re.IGNORECASE,
)
_PDF_FORMULA_CONTEXT_RE = re.compile(
    r"(?:\b(?:equation|formula|where|defined\s+as|expressed\s+as|given\s+by|"
    r"written\s+as|as\s+follows)\b|公式|方程|下式|如下式|表达式|式中|表示为|定义为|"
    r"计算(?:可)?得|计算获得|写为)",
    re.IGNORECASE,
)
_PDF_FORMULA_WORD_RE = re.compile(r"[A-Za-z]{2,}")
_PDF_FORMULA_SINGLE_VARIABLE_RE = re.compile(
    r"(?<![A-Za-z])[A-Za-z](?![A-Za-z])|[\u0370-\u03ff\u1f00-\u1fff]"
)
_PDF_FORMULA_NUMBER_RE = re.compile(r"(?<![A-Za-z])\d+(?:\.\d+)?")
_PDF_FORMULA_MATH_WORDS = frozenset(
    {"sin", "cos", "tan", "log", "ln", "exp", "max", "min", "lim", "det", "arg", "mod"}
)
_PDF_FORMULA_MATH_FONT_TOKENS = (
    "math",
    "symbol",
    "cambria math",
    "stix",
    "cmmi",
    "cmsy",
    "msam",
    "msbm",
    "euclid",
    "mt extra",
)
_PDF_FORMULA_SUPERSCRIPT_TRANSLATION = str.maketrans(
    "⁰¹²³⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉",
    "01234567890123456789",
)
_pdf_formula_candidate_cache: OrderedDict[str, tuple["PdfFormulaCandidate", ...]] = OrderedDict()
_pdf_formula_candidate_cache_lock = Lock()


@dataclass(frozen=True)
class ExtractedDocumentPayload:
    """Structured result of document extraction.

    Default PyMuPDF path preserves the legacy text/blocks/markdown values and
    additionally carries bounded parser provenance for downstream revision
    accounting. Non-PDF extraction leaves ``parser_provenance`` as None.

    Optional external parser paths may add ``blocks`` (structured PDF blocks)
    and ``markdown_full`` (full-document markdown for sidecar writing). Upload
    layer routes these to the chunker (`blocks=`) and the sidecar writer.
    """

    content: str
    blocks: list[StructuredBlock] | None = None  # type: ignore[valid-type]
    markdown_full: str | None = None
    ocr_report: object | None = None
    parser_provenance: PDFParserProvenance | None = None  # type: ignore[valid-type]
    parser_output_sha256: str | None = None


def _document_text_sha256(value: str) -> str:
    """Return the exact UTF-8 text fingerprint used by revision accounting."""

    if not isinstance(value, str):
        raise TypeError("document text must be a string")
    return f"sha256:{hashlib.sha256(value.encode('utf-8')).hexdigest()}"


@dataclass(frozen=True)
class PdfFormulaCandidate:
    """One whole-formula selection target in displayed PDF coordinates.

    Attributes:
        candidate_id: Stable identifier derived from source geometry/text, or
            the persisted formula chunk id when one already exists.
        page: One-based PDF page number.
        bbox: Normalized ``(x, y, width, height)`` in the rotated display page.
        text: Bounded extracted formula text when a text layer is available.
        chunk_id: Best-effort material chunk association.
    """

    candidate_id: str
    page: int
    bbox: tuple[float, float, float, float]
    text: str | None = None
    chunk_id: str | None = None


@dataclass(frozen=True)
class _PdfFormulaLine:
    """Internal PyMuPDF line that passed the conservative formula gate."""

    block_index: int
    line_index: int
    text: str
    rect: tuple[float, float, float, float]
    font_size: float


@dataclass(frozen=True)
class _PdfFormulaFragment:
    """Compact text block used only to recover spatially split formulas."""

    block_index: int
    text: str
    rect: tuple[float, float, float, float]
    font_size: float
    math_font: bool
    scripted: bool


@dataclass(frozen=True)
class _PdfVisualImage:
    """Internal image block used to bind embedded PDF pixels to captions."""

    page: int
    bbox: list[float]
    path: str
    data: bytes
    center_y: float


@dataclass(frozen=True)
class _PdfVisualRegion:
    """Internal rendered region used when PDF visuals are vector/text layout."""

    page: int
    bbox: list[float]
    kind: str


@dataclass(frozen=True)
class _PdfFigureTableRef:
    """Stable figure/table identifier parsed from captions or body mentions."""

    kind: str
    number: str
    ref_id: str


@dataclass(frozen=True)
class _BrowserImageAsset:
    """Browser-displayable bytes and extension for an extracted PDF image."""

    data: bytes
    extension: str


def _safe_asset_segment(value: str) -> str:
    """Return a path segment suitable for project-local extracted assets."""

    normalized = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value or "").strip())
    return normalized.strip("._")[:80] or "document"


def _normalize_image_extension(value: Any) -> str:
    """Return a lowercase extension token without a leading dot."""

    ext = str(value or "").lower().strip().lstrip(".")
    return _IMAGE_EXTENSION_ALIASES.get(ext, ext)


def _normal_visual_number(value: str) -> str:
    """Normalize figure/table numbering while preserving author notation."""

    normalized = str(value or "").strip().replace("–", "-").replace("—", "-")
    return normalized.rstrip(".,:;；。").strip()


def _visual_ref_kind(prefix: str) -> str:
    """Return ``figure`` or ``table`` for a parsed caption/reference prefix."""

    lowered = str(prefix or "").strip().lower().rstrip(".")
    return "table" if lowered in {"表", "table"} else "figure"


def _visual_ref_id(kind: str, source_digest: str, number: str) -> str:
    """Build a stable document-scoped visual evidence id."""

    safe_number = re.sub(r"[^A-Za-z0-9_.-]+", "_", _normal_visual_number(number)).strip("._")
    safe_number = (safe_number or "x")[:40].lower()
    return f"{kind}:{source_digest}:{safe_number}"


def _ref_from_match(match: re.Match[str], source_digest: str) -> _PdfFigureTableRef:
    """Convert a regex match into a stable figure/table reference."""

    kind = _visual_ref_kind(match.group("prefix"))
    number = _normal_visual_number(match.group("number"))
    return _PdfFigureTableRef(
        kind=kind,
        number=number,
        ref_id=_visual_ref_id(kind, source_digest, number),
    )


def _dedupe_refs(refs: list[_PdfFigureTableRef], *, kind: str) -> list[str]:
    """Return stable ids for refs of a kind while preserving first mention."""

    output: list[str] = []
    seen: set[str] = set()
    for ref in refs:
        if ref.kind != kind or ref.ref_id in seen:
            continue
        seen.add(ref.ref_id)
        output.append(ref.ref_id)
    return output


def _figure_table_refs(text: str, source_digest: str) -> list[_PdfFigureTableRef]:
    """Parse figure/table references from one block of text."""

    refs: list[_PdfFigureTableRef] = []
    seen: set[tuple[str, str]] = set()
    for match in _PDF_FIGURE_TABLE_REF_RE.finditer(text or ""):
        ref = _ref_from_match(match, source_digest)
        key = (ref.kind, ref.number.lower())
        if key in seen:
            continue
        seen.add(key)
        refs.append(ref)
    return refs


def _caption_block_type(kind: str) -> str:
    """Return the StructuredBlock type used for a visual caption."""

    return "TableCaption" if kind == "table" else "FigureCaption"


def _convert_image_bytes_to_png(image_bytes: bytes) -> bytes | None:
    """Transcode arbitrary embedded image bytes to PNG for browser preview."""

    if not image_bytes:
        return None
    try:
        from PIL import Image

        with Image.open(io.BytesIO(image_bytes)) as image:
            if image.mode not in {"RGB", "RGBA", "L"}:
                image = image.convert("RGBA" if "A" in image.getbands() else "RGB")
            output = io.BytesIO()
            image.save(output, format="PNG")
            return output.getvalue()
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        _LOGGER.warning("pdf_visual_image_transcode_failed err=%s", exc)
        return None


def _browser_image_asset(image_bytes: bytes | bytearray, raw_ext: Any) -> _BrowserImageAsset | None:
    """Return browser-safe image bytes for a PyMuPDF image block.

    Args:
        image_bytes: Raw embedded PDF image bytes.
        raw_ext: Format token reported by PyMuPDF.

    Returns:
        Browser-displayable bytes plus an extension, or ``None`` when the image
        cannot be converted without producing a broken preview link.
    """

    if not isinstance(image_bytes, (bytes, bytearray)) or not image_bytes:
        return None
    image_data = bytes(image_bytes)
    ext = _normalize_image_extension(raw_ext)
    if ext in _BROWSER_SAFE_IMAGE_EXTENSIONS:
        return _BrowserImageAsset(data=image_data, extension=ext)
    converted = _convert_image_bytes_to_png(image_data)
    if converted is None:
        _LOGGER.warning("pdf_visual_image_unsupported_format ext=%s", ext or "<empty>")
        return None
    return _BrowserImageAsset(data=converted, extension="png")


def _normalized_bbox(rect: Any, page_rect: Any) -> list[float] | None:
    """Convert a PyMuPDF rectangle to normalized [x, y, width, height]."""

    try:
        page_width = float(getattr(page_rect, "width", 0.0) or 0.0)
        page_height = float(getattr(page_rect, "height", 0.0) or 0.0)
        if page_width <= 0.0 or page_height <= 0.0:
            return None
        x0 = max(0.0, float(rect[0])) / page_width
        y0 = max(0.0, float(rect[1])) / page_height
        x1 = min(page_width, float(rect[2])) / page_width
        y1 = min(page_height, float(rect[3])) / page_height
        width = max(0.0, x1 - x0)
        height = max(0.0, y1 - y0)
    except (TypeError, ValueError, AttributeError, IndexError):
        return None
    if width <= 0.0 or height <= 0.0:
        return None
    return [round(x0, 6), round(y0, 6), round(width, 6), round(height, 6)]


def _text_from_pymupdf_block(block: dict[str, Any]) -> str:
    """Extract readable text from a PyMuPDF dict text block."""

    lines: list[str] = []
    raw_lines = block.get("lines")
    if not isinstance(raw_lines, list):
        return ""
    for line in raw_lines:
        if not isinstance(line, dict):
            continue
        spans = line.get("spans")
        if not isinstance(spans, list):
            continue
        line_text = "".join(str(span.get("text") or "") for span in spans if isinstance(span, dict)).strip()
        if line_text:
            lines.append(line_text)
    return "\n".join(lines).strip()


def _normalize_formula_text(value: object, *, max_chars: int = _PDF_FORMULA_TEXT_MAX_CHARS) -> str:
    """Return bounded one-line formula text without invisible separators."""

    if not isinstance(max_chars, int) or isinstance(max_chars, bool) or max_chars < 1:
        raise ValueError("max_chars must be a positive integer")
    if not isinstance(value, str):
        return ""
    normalized = value.replace("\u200b", "").replace("\ufeff", "")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized[:max_chars].rstrip()


def _normalize_formula_fragment_text(
    value: object,
    *,
    max_chars: int = _PDF_FORMULA_TEXT_MAX_CHARS,
) -> str:
    """Normalize full-width OCR fragments without widening the line scanner."""

    if not isinstance(value, str):
        return ""
    normalized = unicodedata.normalize("NFKC", value)
    normalized = "".join(
        " " if unicodedata.category(character) in {"Cf", "Co"} else character
        for character in normalized
    )
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized[:max_chars].rstrip()


def _coerce_formula_rect(value: object) -> tuple[float, float, float, float] | None:
    """Validate a PyMuPDF ``(x0, y0, x1, y1)`` rectangle."""

    if not isinstance(value, Sequence) or isinstance(value, (str, bytes)) or len(value) != 4:
        return None
    coordinates: list[float] = []
    for item in value:
        if isinstance(item, bool) or not isinstance(item, (int, float)):
            return None
        coordinate = float(item)
        if not math.isfinite(coordinate):
            return None
        coordinates.append(coordinate)
    x0, y0, x1, y1 = coordinates
    if x1 <= x0 or y1 <= y0:
        return None
    return (x0, y0, x1, y1)


def _formula_line_payload(raw_line: Mapping[str, Any]) -> tuple[str, bool, bool, float]:
    """Read text plus math-font/script hints from one PyMuPDF line mapping."""

    spans = raw_line.get("spans")
    if not isinstance(spans, list):
        return "", False, False, 0.0
    parts: list[str] = []
    font_sizes: list[float] = []
    baselines: list[float] = []
    math_font = False
    scripted = False
    for span in spans:
        if not isinstance(span, Mapping):
            continue
        text = span.get("text")
        if isinstance(text, str):
            parts.append(text)
            scripted = scripted or _PDF_FORMULA_SCRIPT_RE.search(text) is not None
        raw_size = span.get("size")
        if isinstance(raw_size, (int, float)) and not isinstance(raw_size, bool):
            size = float(raw_size)
            if math.isfinite(size) and size > 0.0:
                font_sizes.append(size)
        raw_flags = span.get("flags")
        if isinstance(raw_flags, int) and not isinstance(raw_flags, bool):
            scripted = scripted or bool(raw_flags & 1)
        font_name = str(span.get("font") or "").casefold()
        math_font = math_font or any(token in font_name for token in _PDF_FORMULA_MATH_FONT_TOKENS)
        origin = span.get("origin")
        if (
            isinstance(origin, Sequence)
            and not isinstance(origin, (str, bytes))
            and len(origin) >= 2
            and isinstance(origin[1], (int, float))
            and not isinstance(origin[1], bool)
        ):
            baseline = float(origin[1])
            if math.isfinite(baseline):
                baselines.append(baseline)

    text = _normalize_formula_text("".join(parts), max_chars=_PDF_FORMULA_TEXT_MAX_CHARS)
    font_size = max(font_sizes, default=0.0)
    if len(font_sizes) >= 2 and max(font_sizes) - min(font_sizes) >= max(0.8, font_size * 0.14):
        scripted = True
    if len(baselines) >= 2 and max(baselines) - min(baselines) >= max(0.8, font_size * 0.12):
        scripted = True
    return text, math_font, scripted, font_size


def _formula_block_fragment(
    raw_block: Mapping[str, Any],
    *,
    block_index: int,
) -> _PdfFormulaFragment | None:
    """Return visible text geometry for one PyMuPDF text block.

    Private-use OCR layout glyphs are intentionally excluded from both text and
    geometry. Some scanned PDFs attach those glyphs to oversized hidden spans,
    which would otherwise make a compact displayed equation look like a broad
    paragraph block.
    """

    raw_lines = raw_block.get("lines")
    if not isinstance(raw_lines, list):
        return None
    line_texts: list[str] = []
    rects: list[tuple[float, float, float, float]] = []
    font_sizes: list[float] = []
    baselines: list[float] = []
    math_font = False
    scripted = False
    for raw_line in raw_lines:
        if not isinstance(raw_line, Mapping):
            continue
        raw_spans = raw_line.get("spans")
        if not isinstance(raw_spans, list):
            continue
        line_parts: list[str] = []
        for raw_span in raw_spans:
            if not isinstance(raw_span, Mapping):
                continue
            text = _normalize_formula_fragment_text(raw_span.get("text"))
            rect = _coerce_formula_rect(raw_span.get("bbox"))
            if not text or rect is None:
                continue
            line_parts.append(text)
            rects.append(rect)
            scripted = scripted or _PDF_FORMULA_SCRIPT_RE.search(text) is not None
            raw_size = raw_span.get("size")
            if isinstance(raw_size, (int, float)) and not isinstance(raw_size, bool):
                size = float(raw_size)
                if math.isfinite(size) and size > 0.0:
                    font_sizes.append(size)
            raw_flags = raw_span.get("flags")
            if isinstance(raw_flags, int) and not isinstance(raw_flags, bool):
                scripted = scripted or bool(raw_flags & 1)
            font_name = str(raw_span.get("font") or "").casefold()
            math_font = math_font or any(
                token in font_name for token in _PDF_FORMULA_MATH_FONT_TOKENS
            )
            origin = raw_span.get("origin")
            if (
                isinstance(origin, Sequence)
                and not isinstance(origin, (str, bytes))
                and len(origin) >= 2
                and isinstance(origin[1], (int, float))
                and not isinstance(origin[1], bool)
            ):
                baseline = float(origin[1])
                if math.isfinite(baseline):
                    baselines.append(baseline)
        line_text = _normalize_formula_text("".join(line_parts))
        if line_text:
            line_texts.append(line_text)
    if not line_texts or not rects:
        return None
    font_size = max(font_sizes, default=0.0)
    if len(font_sizes) >= 2 and max(font_sizes) - min(font_sizes) >= max(0.8, font_size * 0.14):
        scripted = True
    if len(baselines) >= 2 and max(baselines) - min(baselines) >= max(0.8, font_size * 0.12):
        scripted = True
    rect = (
        min(item[0] for item in rects),
        min(item[1] for item in rects),
        max(item[2] for item in rects),
        max(item[3] for item in rects),
    )
    fragment_text = _normalize_formula_text(" ".join(line_texts))
    if len(line_texts) > 1 and len(_PDF_FORMULA_RELATION_RE.findall(fragment_text)) > 1:
        return None
    return _PdfFormulaFragment(
        block_index=block_index,
        text=fragment_text,
        rect=rect,
        font_size=font_size,
        math_font=math_font,
        scripted=scripted,
    )


def _formula_prose_word_count(text: str) -> int:
    """Count natural-language-looking ASCII words in a candidate line."""

    count = 0
    for match in _PDF_FORMULA_WORD_RE.finditer(text):
        token = match.group(0)
        lowered = token.casefold()
        if lowered in _PDF_FORMULA_MATH_WORDS:
            continue
        if token.isupper() and len(token) <= 4:
            continue
        count += 1
    return count


def _formula_line_score(text: str, *, math_font: bool, scripted: bool) -> int | None:
    """Score a line only when conservative structural formula gates pass."""

    normalized = _normalize_formula_text(text, max_chars=_PDF_FORMULA_TEXT_MAX_CHARS)
    if len(normalized) < 3 or len(normalized) > 260:
        return None
    if _PDF_FORMULA_NUMBER_ONLY_RE.fullmatch(normalized):
        return None
    if _PDF_FORMULA_URL_RE.search(normalized):
        return None
    if _PDF_FORMULA_CAPTION_RE.match(normalized) or _PDF_FORMULA_LIST_RE.match(normalized):
        return None
    if _PDF_FORMULA_DISCOURSE_RE.match(normalized):
        return None

    core = _PDF_FORMULA_TRAILING_NUMBER_RE.sub("", normalized).strip()
    relation_count = len(_PDF_FORMULA_RELATION_RE.findall(core))
    strong_math_count = len(_PDF_FORMULA_STRONG_MATH_RE.findall(core))
    if relation_count == 0 and strong_math_count == 0:
        return None

    operator_count = len(_PDF_FORMULA_OPERATOR_RE.findall(core))
    greek_count = len(_PDF_FORMULA_GREEK_RE.findall(core))
    script_count = len(_PDF_FORMULA_SCRIPT_RE.findall(core))
    variable_count = len(_PDF_FORMULA_SINGLE_VARIABLE_RE.findall(core))
    number_count = len(_PDF_FORMULA_NUMBER_RE.findall(core))
    math_word_count = sum(
        1
        for token in _PDF_FORMULA_WORD_RE.findall(core)
        if token.casefold() in _PDF_FORMULA_MATH_WORDS
    )
    prose_word_count = _formula_prose_word_count(core)

    if prose_word_count >= 6:
        return None
    if re.search(r"[\u4e00-\u9fff]{8,}", core) and strong_math_count < 2:
        return None

    operand_count = variable_count + number_count
    secondary_count = operator_count + greek_count + script_count + math_word_count
    secondary_count += int(math_font) + int(scripted)
    if operand_count >= 2 or (strong_math_count > 0 and operand_count >= 1):
        secondary_count += 1
    if secondary_count < 1:
        return None

    # Relation-only prose such as "The setting x = 3 was used" is the main
    # false-positive class. A real displayed formula with prose identifiers
    # must carry stronger math structure than the relation sign itself.
    if prose_word_count >= 2 and strong_math_count == 0 and not scripted:
        return None
    if (
        prose_word_count >= 1
        and strong_math_count == 0
        and operator_count == 0
        and not scripted
        and not math_font
        and variable_count < 2
    ):
        return None

    return (
        relation_count * 5
        + strong_math_count * 6
        + min(secondary_count, 8) * 2
        - min(prose_word_count, 5)
    )


def _formula_fragment_is_compact(
    fragment: _PdfFormulaFragment,
    *,
    page_width: float,
    page_height: float,
) -> bool:
    """Reject paragraph, caption, and degenerate blocks before clustering."""

    if page_width <= 0.0 or page_height <= 0.0:
        return False
    width_ratio = (fragment.rect[2] - fragment.rect[0]) / page_width
    height_ratio = (fragment.rect[3] - fragment.rect[1]) / page_height
    text = fragment.text
    if not (0.003 <= width_ratio <= 0.55 and 0.003 <= height_ratio <= 0.12):
        return False
    if not text or len(text) > 260:
        return False
    if _PDF_FORMULA_URL_RE.search(text):
        return False
    if _PDF_FORMULA_CAPTION_RE.match(text) or _PDF_FORMULA_LIST_RE.match(text):
        return False
    if _PDF_FORMULA_DISCOURSE_RE.match(text):
        return False
    if _formula_prose_word_count(text) > 1:
        return False
    return len(re.findall(r"[\u4e00-\u9fff]", text)) <= 3


def _formula_fragment_is_anchor(fragment: _PdfFormulaFragment) -> bool:
    """Return whether a compact fragment can seed one relation formula."""

    core = _PDF_FORMULA_TRAILING_NUMBER_RE.sub("", fragment.text).strip()
    if _PDF_FORMULA_RELATION_RE.search(core) is None:
        return False
    # OCR may isolate the relation glyph from both operands. Only an exact
    # relation token can seed that recovery; operators such as ± or × remain
    # non-anchors and still require a stronger formula fragment.
    if _PDF_FORMULA_RELATION_RE.fullmatch(core):
        return True
    operand_count = len(_PDF_FORMULA_SINGLE_VARIABLE_RE.findall(core))
    operand_count += len(_PDF_FORMULA_NUMBER_RE.findall(core))
    operand_count += len(_PDF_FORMULA_GREEK_RE.findall(core))
    operand_count += len(_PDF_FORMULA_SCRIPT_RE.findall(core))
    return operand_count >= 1


def _formula_fragment_union_rect(
    fragments: Sequence[_PdfFormulaFragment],
) -> tuple[float, float, float, float]:
    """Return the smallest raw-coordinate rectangle containing fragments."""

    return (
        min(fragment.rect[0] for fragment in fragments),
        min(fragment.rect[1] for fragment in fragments),
        max(fragment.rect[2] for fragment in fragments),
        max(fragment.rect[3] for fragment in fragments),
    )


def _formula_fragment_is_neighbor(
    cluster_rect: tuple[float, float, float, float],
    fragment_rect: tuple[float, float, float, float],
    *,
    page_width: float,
) -> bool:
    """Return whether two compact regions can belong to one displayed line."""

    cluster_height = cluster_rect[3] - cluster_rect[1]
    fragment_height = fragment_rect[3] - fragment_rect[1]
    smaller_height = min(cluster_height, fragment_height)
    vertical_overlap = max(
        0.0,
        min(cluster_rect[3], fragment_rect[3]) - max(cluster_rect[1], fragment_rect[1]),
    )
    center_distance = abs(
        (cluster_rect[1] + cluster_rect[3]) / 2.0
        - (fragment_rect[1] + fragment_rect[3]) / 2.0
    )
    if cluster_rect[2] < fragment_rect[0]:
        horizontal_gap = fragment_rect[0] - cluster_rect[2]
    elif fragment_rect[2] < cluster_rect[0]:
        horizontal_gap = cluster_rect[0] - fragment_rect[2]
    else:
        horizontal_gap = 0.0
    vertically_aligned = vertical_overlap >= smaller_height * 0.15
    vertically_aligned = vertically_aligned or center_distance <= max(6.0, smaller_height * 1.5)
    return vertically_aligned and horizontal_gap <= page_width * 0.065


def _formula_fragment_anchor_owner(
    fragment: _PdfFormulaFragment,
    anchors: Sequence[_PdfFormulaFragment],
) -> int | None:
    """Return the nearest relation anchor for one compact formula fragment."""

    if not anchors:
        return None
    fragment_center_x = (fragment.rect[0] + fragment.rect[2]) / 2.0
    fragment_center_y = (fragment.rect[1] + fragment.rect[3]) / 2.0
    ranked = sorted(
        (
            abs(fragment_center_x - (anchor.rect[0] + anchor.rect[2]) / 2.0)
            + 2.0 * abs(fragment_center_y - (anchor.rect[1] + anchor.rect[3]) / 2.0),
            anchor.block_index,
        )
        for anchor in anchors
    )
    return ranked[0][1]


def _formula_fragment_has_context_support(
    formula_rect: tuple[float, float, float, float],
    fragments: Sequence[_PdfFormulaFragment],
    *,
    member_indexes: set[int],
    page_width: float,
    page_height: float,
) -> bool:
    """Require a nearby equation number or explicit prose cue for fallback."""

    formula_center_y = (formula_rect[1] + formula_rect[3]) / 2.0
    for fragment in fragments:
        if fragment.block_index in member_indexes:
            if _PDF_FORMULA_TRAILING_NUMBER_RE.search(fragment.text):
                return True
            continue
        fragment_center_y = (fragment.rect[1] + fragment.rect[3]) / 2.0
        if _PDF_FORMULA_NUMBER_ONLY_RE.fullmatch(fragment.text):
            horizontal_gap = fragment.rect[0] - formula_rect[2]
            if (
                -2.0 <= horizontal_gap <= page_width * 0.45
                and abs(formula_center_y - fragment_center_y)
                <= max(12.0, formula_rect[3] - formula_rect[1])
            ):
                return True
        if _PDF_FORMULA_CONTEXT_RE.search(fragment.text) is None:
            continue
        if fragment.rect[3] < formula_rect[1]:
            vertical_gap = formula_rect[1] - fragment.rect[3]
        elif formula_rect[3] < fragment.rect[1]:
            vertical_gap = fragment.rect[1] - formula_rect[3]
        else:
            vertical_gap = 0.0
        if vertical_gap <= max(36.0, page_height * 0.09):
            return True
    return False


def _fragmented_formula_lines(
    raw_blocks: Sequence[object],
    *,
    page_width: float,
    page_height: float,
) -> list[_PdfFormulaLine]:
    """Recover displayed formulas split across neighboring OCR text blocks."""

    all_fragments: list[_PdfFormulaFragment] = []
    for block_index, raw_block in enumerate(raw_blocks):
        if not isinstance(raw_block, Mapping) or raw_block.get("type") != 0:
            continue
        fragment = _formula_block_fragment(raw_block, block_index=block_index)
        if fragment is not None:
            all_fragments.append(fragment)
    compact_fragments = [
        fragment
        for fragment in all_fragments
        if _formula_fragment_is_compact(
            fragment,
            page_width=page_width,
            page_height=page_height,
        )
    ]
    if len(compact_fragments) > _PDF_FORMULA_FRAGMENT_MAX_PER_PAGE:
        return []
    anchors = [fragment for fragment in compact_fragments if _formula_fragment_is_anchor(fragment)]
    if len(anchors) > _PDF_FORMULA_ANCHOR_MAX_PER_PAGE:
        return []
    output: list[_PdfFormulaLine] = []
    seen_member_sets: set[tuple[int, ...]] = set()
    for anchor in anchors:
        members = [anchor]
        member_indexes = {anchor.block_index}
        changed = True
        while changed:
            changed = False
            cluster_rect = _formula_fragment_union_rect(members)
            for fragment in compact_fragments:
                if fragment.block_index in member_indexes:
                    continue
                # Keep standalone equation numbers outside the recovered line:
                # they are context evidence here and are paired exactly once
                # by the page-level number matcher below.
                if _PDF_FORMULA_NUMBER_ONLY_RE.fullmatch(fragment.text):
                    continue
                # A second relation-bearing fragment is a separate formula
                # anchor unless it is part of the same OCR block. Do not let a
                # greedy spatial walk merge two neighboring equations.
                if (
                    fragment.block_index != anchor.block_index
                    and _formula_fragment_is_anchor(fragment)
                ):
                    continue
                if _formula_fragment_anchor_owner(fragment, anchors) != anchor.block_index:
                    continue
                if not _formula_fragment_is_neighbor(
                    cluster_rect,
                    fragment.rect,
                    page_width=page_width,
                ):
                    continue
                trial_members = [*members, fragment]
                trial_rect = _formula_fragment_union_rect(trial_members)
                width_ratio = (trial_rect[2] - trial_rect[0]) / page_width
                height_ratio = (trial_rect[3] - trial_rect[1]) / page_height
                if width_ratio > 0.75 or height_ratio > 0.12:
                    continue
                members.append(fragment)
                member_indexes.add(fragment.block_index)
                changed = True
        if len(members) < 2:
            continue
        member_key = tuple(sorted(member_indexes))
        if member_key in seen_member_sets:
            continue
        seen_member_sets.add(member_key)
        members.sort(
            key=lambda fragment: (
                fragment.rect[0],
                fragment.rect[1],
                fragment.block_index,
            )
        )
        text = _normalize_formula_text(" ".join(fragment.text for fragment in members))
        if len(re.findall(r"[\u4e00-\u9fff]", text)) > 3:
            continue
        math_font = any(fragment.math_font for fragment in members)
        scripted = any(fragment.scripted for fragment in members)
        if _formula_line_score(text, math_font=math_font, scripted=scripted) is None:
            continue
        rect = _formula_fragment_union_rect(members)
        if not _formula_line_geometry_is_local(
            rect,
            page_width=page_width,
            page_height=page_height,
        ):
            continue
        if not _formula_fragment_has_context_support(
            rect,
            all_fragments,
            member_indexes=member_indexes,
            page_width=page_width,
            page_height=page_height,
        ):
            continue
        font_size = max((fragment.font_size for fragment in members), default=0.0)
        vertical_margin = max(1.5, min(3.0, font_size * 0.25))
        output.append(
            _PdfFormulaLine(
                block_index=min(member_indexes),
                line_index=-1,
                text=text,
                rect=(rect[0], rect[1] - vertical_margin, rect[2], rect[3] + vertical_margin),
                font_size=font_size,
            )
        )
    return output


def _formula_raw_rect_overlap_ratio(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    """Return intersection area relative to the smaller raw PDF rectangle."""

    overlap_width = max(0.0, min(first[2], second[2]) - max(first[0], second[0]))
    overlap_height = max(0.0, min(first[3], second[3]) - max(first[1], second[1]))
    first_area = (first[2] - first[0]) * (first[3] - first[1])
    second_area = (second[2] - second[0]) * (second[3] - second[1])
    return (overlap_width * overlap_height) / max(min(first_area, second_area), 1e-9)


def _formula_line_geometry_is_local(
    rect: tuple[float, float, float, float],
    *,
    page_width: float,
    page_height: float,
) -> bool:
    """Reject full-width prose blocks and degenerate glyph fragments."""

    if page_width <= 0.0 or page_height <= 0.0:
        return False
    width_ratio = (rect[2] - rect[0]) / page_width
    height_ratio = (rect[3] - rect[1]) / page_height
    return 0.03 <= width_ratio <= 0.96 and 0.006 <= height_ratio <= 0.20


def _formula_number_geometry_is_local(
    rect: tuple[float, float, float, float],
    *,
    page_width: float,
    page_height: float,
) -> bool:
    """Allow a narrow equation number only for later formula-line pairing."""

    if page_width <= 0.0 or page_height <= 0.0:
        return False
    width_ratio = (rect[2] - rect[0]) / page_width
    height_ratio = (rect[3] - rect[1]) / page_height
    return 0.005 <= width_ratio <= 0.20 and 0.006 <= height_ratio <= 0.20


def _formula_number_for_line(
    formula: _PdfFormulaLine,
    formula_lines: Sequence[_PdfFormulaLine],
    number_lines: Sequence[_PdfFormulaLine],
    used_numbers: set[tuple[int, int]],
    *,
    page_width: float,
) -> _PdfFormulaLine | None:
    """Return a nearby standalone equation number, never a page number alone."""

    if _PDF_FORMULA_TRAILING_NUMBER_RE.search(formula.text):
        return None
    formula_center_y = (formula.rect[1] + formula.rect[3]) / 2.0
    tolerance = max(6.0, formula.font_size * 0.8)
    best: tuple[float, _PdfFormulaLine] | None = None
    for number_line in number_lines:
        key = (number_line.block_index, number_line.line_index)
        if key in used_numbers:
            continue
        number_center_y = (number_line.rect[1] + number_line.rect[3]) / 2.0
        vertical_distance = abs(formula_center_y - number_center_y)
        vertical_overlap = max(
            0.0,
            min(formula.rect[3], number_line.rect[3]) - max(formula.rect[1], number_line.rect[1]),
        )
        if vertical_distance > tolerance and vertical_overlap <= 0.0:
            continue
        if formula.rect[2] < number_line.rect[0]:
            horizontal_gap = number_line.rect[0] - formula.rect[2]
        elif number_line.rect[2] < formula.rect[0]:
            horizontal_gap = formula.rect[0] - number_line.rect[2]
        else:
            horizontal_gap = 0.0
        if page_width <= 0.0 or horizontal_gap > page_width * 0.45:
            continue
        owner_key = min(
            (
                max(
                    0.0,
                    max(owner.rect[0], number_line.rect[0])
                    - min(owner.rect[2], number_line.rect[2]),
                ),
                owner.block_index,
                owner.line_index,
            )
            for owner in formula_lines
            if (
                abs(
                    (owner.rect[1] + owner.rect[3]) / 2.0
                    - number_center_y
                )
                <= max(6.0, owner.font_size * 0.8)
                or min(owner.rect[3], number_line.rect[3])
                > max(owner.rect[1], number_line.rect[1])
            )
        )
        if owner_key[1:] != (formula.block_index, formula.line_index):
            continue
        same_block_penalty = 0.0 if formula.block_index == number_line.block_index else 1.5
        side_penalty = 0.0 if number_line.rect[0] >= formula.rect[2] - 2.0 else 1.0
        score = vertical_distance + same_block_penalty + side_penalty
        if best is None or score < best[0]:
            best = (score, number_line)
    return best[1] if best is not None else None


def _union_formula_rect(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float] | None,
) -> tuple[float, float, float, float]:
    """Return a raw-coordinate rectangle union."""

    if second is None:
        return first
    return (
        min(first[0], second[0]),
        min(first[1], second[1]),
        max(first[2], second[2]),
        max(first[3], second[3]),
    )


def _normalized_formula_display_bbox(
    raw_rect: tuple[float, float, float, float],
    page: Any,
) -> tuple[float, float, float, float] | None:
    """Rotate a raw text bbox into the displayed page before normalization."""

    try:
        import pymupdf

        page_rect = new_pymupdf_rect(pymupdf, page.rect)
        display_rect = transform_pymupdf_rect(
            new_pymupdf_rect(pymupdf, raw_rect),
            page.rotation_matrix,
        )
        page_width = float(getattr(page_rect, "width"))
        page_height = float(getattr(page_rect, "height"))
        if page_width <= 0.0 or page_height <= 0.0:
            return None
        page_x0 = float(getattr(page_rect, "x0"))
        page_y0 = float(getattr(page_rect, "y0"))
        left = max(page_x0, float(getattr(display_rect, "x0")) - 3.0)
        top = max(page_y0, float(getattr(display_rect, "y0")) - 2.0)
        right = min(float(getattr(page_rect, "x1")), float(getattr(display_rect, "x1")) + 3.0)
        bottom = min(float(getattr(page_rect, "y1")), float(getattr(display_rect, "y1")) + 2.0)
        if right <= left or bottom <= top:
            return None
        normalized = (
            round((left - page_x0) / page_width, 6),
            round((top - page_y0) / page_height, 6),
            round((right - left) / page_width, 6),
            round((bottom - top) / page_height, 6),
        )
    except (ImportError, AttributeError, RuntimeError, TypeError, ValueError):
        return None
    return _coerce_normalized_formula_bbox(normalized)


def _formula_candidate_id(
    *,
    page: int,
    bbox: tuple[float, float, float, float],
    text: str,
) -> str:
    """Build a deterministic id independent from the machine-local file path."""

    fingerprint = f"{page}|{'|'.join(f'{value:.6f}' for value in bbox)}|{_formula_match_key(text)}"
    digest = hashlib.sha256(fingerprint.encode("utf-8", errors="ignore")).hexdigest()[:16]
    return f"formula-p{page:04d}-{digest}"


def _formula_bbox_overlap_ratio(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    """Return intersection area relative to the smaller candidate box."""

    overlap_width = max(
        0.0,
        min(first[0] + first[2], second[0] + second[2]) - max(first[0], second[0]),
    )
    overlap_height = max(
        0.0,
        min(first[1] + first[3], second[1] + second[3]) - max(first[1], second[1]),
    )
    smaller_area = min(first[2] * first[3], second[2] * second[3])
    return (overlap_width * overlap_height) / max(smaller_area, 1e-9)


def _formula_bbox_iou(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    """Return intersection-over-union for two normalized candidate boxes."""

    overlap_width = max(
        0.0,
        min(first[0] + first[2], second[0] + second[2]) - max(first[0], second[0]),
    )
    overlap_height = max(
        0.0,
        min(first[1] + first[3], second[1] + second[3]) - max(first[1], second[1]),
    )
    intersection = overlap_width * overlap_height
    union = first[2] * first[3] + second[2] * second[3] - intersection
    return intersection / max(union, 1e-9)


def _formula_match_key(value: object) -> str:
    """Normalize visible and LaTeX-like formula text for dedupe/chunk joins."""

    normalized = _normalize_formula_text(value).translate(_PDF_FORMULA_SUPERSCRIPT_TRANSLATION)
    normalized = normalized.casefold().replace("−", "-").replace("×", "*").replace("÷", "/")
    return re.sub(r"[^0-9a-z\u0370-\u03ff=<>+*/.-]+", "", normalized)


def _formula_candidates_are_duplicates(
    first: PdfFormulaCandidate,
    second: PdfFormulaCandidate,
) -> bool:
    """Detect duplicate hidden text layers without collapsing repeated equations."""

    if first.page != second.page:
        return False
    geometry_match = _formula_bbox_overlap_ratio(first.bbox, second.bbox) >= 0.78
    first_text = _formula_match_key(first.text)
    second_text = _formula_match_key(second.text)
    if (
        _formula_bbox_iou(first.bbox, second.bbox) >= 0.70
        and first.chunk_id
        and first.chunk_id == second.chunk_id
    ):
        return True
    if geometry_match and (not first_text or not second_text or first_text == second_text):
        return True
    if first_text and first_text == second_text:
        first_center_x = first.bbox[0] + first.bbox[2] / 2.0
        second_center_x = second.bbox[0] + second.bbox[2] / 2.0
        first_center_y = first.bbox[1] + first.bbox[3] / 2.0
        second_center_y = second.bbox[1] + second.bbox[3] / 2.0
        horizontal_tolerance = max(0.01, min(first.bbox[2], second.bbox[2]) * 0.25)
        return (
            abs(first_center_x - second_center_x) <= horizontal_tolerance
            and abs(first_center_y - second_center_y) <= 0.004
        )
    return False


def _formula_candidate_has_atomic_replacements(
    candidate: PdfFormulaCandidate,
    detected: Sequence[PdfFormulaCandidate],
) -> bool:
    """Return whether two distinct local boxes safely supersede one broad chunk box."""

    if not candidate.chunk_id:
        return False
    candidate_area = candidate.bbox[2] * candidate.bbox[3]
    if candidate_area <= 0.0:
        return False
    replacements: list[tuple[PdfFormulaCandidate, float]] = []
    for replacement in detected:
        if (
            replacement.page != candidate.page
            or replacement.chunk_id != candidate.chunk_id
            or _formula_bbox_overlap_ratio(candidate.bbox, replacement.bbox) < 0.78
        ):
            continue
        replacement_area = replacement.bbox[2] * replacement.bbox[3]
        area_ratio = replacement_area / candidate_area
        if 0.20 <= area_ratio <= 0.75:
            replacements.append((replacement, replacement_area))
    for index, (first, first_area) in enumerate(replacements):
        for second, second_area in replacements[index + 1 :]:
            if first_area + second_area < candidate_area * 0.60:
                continue
            if (
                _formula_bbox_overlap_ratio(first.bbox, second.bbox) <= 0.15
                and not _formula_candidates_are_duplicates(first, second)
            ):
                return True
    return False


def _scan_pymupdf_formula_candidates(source_path: Path) -> tuple[PdfFormulaCandidate, ...]:
    """Scan one PDF text layer, retaining no page/image payload after return."""

    try:
        import pymupdf
    except ImportError:
        return ()

    candidates: list[PdfFormulaCandidate] = []
    try:
        with open_pymupdf_document(pymupdf, str(source_path)) as document:
            text_flags = int(getattr(pymupdf, "TEXTFLAGS_DICT", 0))
            image_flag = int(getattr(pymupdf, "TEXT_PRESERVE_IMAGES", 4))
            text_flags &= ~image_flag
            for page_number, page in enumerate(document, start=1):
                page_dict = page.get_text("dict", sort=True, flags=text_flags)
                if not isinstance(page_dict, Mapping):
                    continue
                raw_blocks = page_dict.get("blocks")
                if not isinstance(raw_blocks, list):
                    continue
                raw_page_width = float(page_dict.get("width") or 0.0)
                raw_page_height = float(page_dict.get("height") or 0.0)
                if raw_page_width <= 0.0 or raw_page_height <= 0.0:
                    continue

                formula_lines: list[_PdfFormulaLine] = []
                number_lines: list[_PdfFormulaLine] = []
                for block_index, raw_block in enumerate(raw_blocks):
                    if not isinstance(raw_block, Mapping) or raw_block.get("type") != 0:
                        continue
                    raw_lines = raw_block.get("lines")
                    if not isinstance(raw_lines, list):
                        continue
                    for line_index, raw_line in enumerate(raw_lines):
                        if not isinstance(raw_line, Mapping):
                            continue
                        rect = _coerce_formula_rect(raw_line.get("bbox"))
                        if rect is None:
                            continue
                        text, math_font, scripted, font_size = _formula_line_payload(raw_line)
                        if not text:
                            continue
                        line = _PdfFormulaLine(
                            block_index=block_index,
                            line_index=line_index,
                            text=text,
                            rect=rect,
                            font_size=font_size,
                        )
                        if _PDF_FORMULA_NUMBER_ONLY_RE.fullmatch(text):
                            if _formula_number_geometry_is_local(
                                rect,
                                page_width=raw_page_width,
                                page_height=raw_page_height,
                            ):
                                number_lines.append(line)
                            continue
                        if not _formula_line_geometry_is_local(
                            rect,
                            page_width=raw_page_width,
                            page_height=raw_page_height,
                        ):
                            continue
                        if (
                            _formula_line_score(text, math_font=math_font, scripted=scripted)
                            is not None
                        ):
                            formula_lines.append(line)

                fragmented_lines = _fragmented_formula_lines(
                    raw_blocks,
                    page_width=raw_page_width,
                    page_height=raw_page_height,
                )
                if fragmented_lines:
                    formula_lines = [
                        line
                        for line in formula_lines
                        if not any(
                            _formula_raw_rect_overlap_ratio(line.rect, fragmented.rect) >= 0.78
                            for fragmented in fragmented_lines
                        )
                    ]
                    formula_lines.extend(fragmented_lines)

                formula_lines.sort(
                    key=lambda item: (item.rect[1], item.rect[0], item.block_index, item.line_index)
                )
                number_lines.sort(
                    key=lambda item: (item.rect[1], item.rect[0], item.block_index, item.line_index)
                )
                used_numbers: set[tuple[int, int]] = set()
                for formula_line in formula_lines:
                    number_line = _formula_number_for_line(
                        formula_line,
                        formula_lines,
                        number_lines,
                        used_numbers,
                        page_width=raw_page_width,
                    )
                    if number_line is not None:
                        used_numbers.add((number_line.block_index, number_line.line_index))
                    raw_bbox = _union_formula_rect(
                        formula_line.rect,
                        number_line.rect if number_line is not None else None,
                    )
                    bbox = _normalized_formula_display_bbox(raw_bbox, page)
                    if bbox is None:
                        continue
                    text = formula_line.text
                    if number_line is not None:
                        text = _normalize_formula_text(f"{text} {number_line.text}")
                    candidate = PdfFormulaCandidate(
                        candidate_id=_formula_candidate_id(page=page_number, bbox=bbox, text=text),
                        page=page_number,
                        bbox=bbox,
                        text=text or None,
                    )
                    if any(
                        _formula_candidates_are_duplicates(existing, candidate)
                        for existing in candidates
                    ):
                        continue
                    candidates.append(candidate)
                    if len(candidates) >= _PDF_FORMULA_SCAN_LIMIT:
                        return tuple(candidates)
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        _LOGGER.warning(
            "pdf_formula_candidate_scan_failed filename=%s err=%s", source_path.name, exc
        )
        return ()
    return tuple(candidates)


def _formula_candidate_cache_key(source_path: Path) -> str | None:
    """Return a cache key tied to source identity and detector semantics."""

    try:
        resolved = source_path.resolve(strict=True)
        stat = resolved.stat()
    except (OSError, RuntimeError, ValueError):
        return None
    return (
        f"{str(resolved).casefold()}|{stat.st_size}|{stat.st_mtime_ns}|"
        f"{_PDF_FORMULA_DETECTOR_VERSION}"
    )


def extract_pymupdf_formula_candidates(
    source_path: Path,
    *,
    limit: int = _PDF_FORMULA_SCAN_LIMIT,
) -> list[PdfFormulaCandidate]:
    """Return bounded whole-formula candidates from a PDF text layer.

    Args:
        source_path: Existing PDF source path resolved by the resources router.
        limit: Maximum candidates to return, from 1 through 200.

    Returns:
        Stable page-ordered candidates in rotated display coordinates. Missing
        files, non-PDF sources, unavailable PyMuPDF, and unreadable PDFs return
        an empty list. The process-local cache stores only small immutable
        candidate tuples and is invalidated by source size/mtime.

    Raises:
        ValueError: If ``limit`` is outside the public endpoint bound.
    """

    if (
        not isinstance(limit, int)
        or isinstance(limit, bool)
        or limit < 1
        or limit > _PDF_FORMULA_SCAN_LIMIT
    ):
        raise ValueError(f"limit must be between 1 and {_PDF_FORMULA_SCAN_LIMIT}")
    if not isinstance(source_path, Path) or source_path.suffix.casefold() != ".pdf":
        return []
    cache_key = _formula_candidate_cache_key(source_path)
    if cache_key is None:
        return []
    with _pdf_formula_candidate_cache_lock:
        cached = _pdf_formula_candidate_cache.get(cache_key)
        if cached is not None:
            _pdf_formula_candidate_cache.move_to_end(cache_key)
            return list(cached[:limit])

    scanned = _scan_pymupdf_formula_candidates(source_path)
    with _pdf_formula_candidate_cache_lock:
        cached = _pdf_formula_candidate_cache.get(cache_key)
        if cached is None:
            _pdf_formula_candidate_cache[cache_key] = scanned
            while len(_pdf_formula_candidate_cache) > _PDF_FORMULA_CACHE_MAX:
                _pdf_formula_candidate_cache.popitem(last=False)
            cached = scanned
        else:
            _pdf_formula_candidate_cache.move_to_end(cache_key)
    return list(cached[:limit])


def _coerce_normalized_formula_bbox(
    value: object,
) -> tuple[float, float, float, float] | None:
    """Validate normalized ``(x, y, width, height)`` selection geometry."""

    if not isinstance(value, Sequence) or isinstance(value, (str, bytes)) or len(value) != 4:
        return None
    coordinates: list[float] = []
    for item in value:
        if isinstance(item, bool) or not isinstance(item, (int, float)):
            return None
        coordinate = float(item)
        if not math.isfinite(coordinate):
            return None
        coordinates.append(coordinate)
    x, y, width, height = coordinates
    tolerance = 1e-6
    if x < 0.0 or y < 0.0 or width <= 0.0 or height <= 0.0:
        return None
    if x + width > 1.0 + tolerance or y + height > 1.0 + tolerance:
        return None
    return tuple(round(max(0.0, min(1.0, coordinate)), 6) for coordinate in coordinates)  # type: ignore[return-value]


def _formula_chunk_text(chunk: Mapping[str, Any]) -> str:
    """Return the highest-signal bounded formula text stored on a chunk."""

    for key in ("equation_latex", "raw_content", "content"):
        text = _normalize_formula_text(chunk.get(key))
        if text:
            if key == "content" and "\n" in str(chunk.get(key)):
                text = _normalize_formula_text(str(chunk.get(key)).rsplit("\n", 1)[-1]) or text
            return text
    return ""


def _formula_candidate_sort_key(candidate: PdfFormulaCandidate) -> tuple[object, ...]:
    """Return deterministic document order for mixed candidate sources."""

    return (
        candidate.page,
        round(candidate.bbox[1], 6),
        round(candidate.bbox[0], 6),
        candidate.candidate_id,
    )


def formula_candidates_from_chunks(
    chunks: Sequence[Mapping[str, Any]],
    *,
    material_id: str,
    limit: int = _PDF_FORMULA_SCAN_LIMIT,
) -> list[PdfFormulaCandidate]:
    """Project reliable persisted formula/equation chunks into UI candidates.

    Args:
        chunks: Already-loaded finite material chunk sequence.
        material_id: Material scope used only for fallback stable ids.
        limit: Maximum candidates to return, from 1 through 200.

    Returns:
        Formula chunks with one-based pages and valid normalized bboxes. Broad
        narrative chunks and unknown coordinate units are excluded.
    """

    if (
        not isinstance(limit, int)
        or isinstance(limit, bool)
        or limit < 1
        or limit > _PDF_FORMULA_SCAN_LIMIT
    ):
        raise ValueError(f"limit must be between 1 and {_PDF_FORMULA_SCAN_LIMIT}")
    normalized_material_id = str(material_id or "").strip()
    if not normalized_material_id:
        raise ValueError("material_id must be a non-empty string")

    candidates: list[PdfFormulaCandidate] = []
    seen_ids: set[str] = set()
    for index, chunk in enumerate(chunks):
        if not isinstance(chunk, Mapping):
            continue
        chunk_material_id = str(chunk.get("material_id") or normalized_material_id).strip()
        if chunk_material_id and chunk_material_id != normalized_material_id:
            continue
        chunk_type = str(chunk.get("chunk_type") or "").strip().casefold()
        equation_latex = _normalize_formula_text(chunk.get("equation_latex"))
        if chunk_type not in {"formula", "equation"} and not equation_latex:
            continue
        raw_page = chunk.get("page")
        if isinstance(raw_page, bool) or not isinstance(raw_page, int) or raw_page < 1:
            continue
        raw_unit = chunk.get("bbox_unit")
        unit = str(getattr(raw_unit, "value", raw_unit) or "").strip().casefold()
        if unit != "normalized_ratio":
            continue
        bbox = _coerce_normalized_formula_bbox(chunk.get("bbox"))
        if bbox is None:
            continue
        chunk_id = str(chunk.get("chunk_id") or "").strip()[:200] or None
        text = equation_latex or _formula_chunk_text(chunk)
        candidate_id = chunk_id
        if candidate_id is None:
            seed_text = f"{normalized_material_id}|{index}|{text}"
            candidate_id = _formula_candidate_id(page=raw_page, bbox=bbox, text=seed_text)
        if candidate_id in seen_ids:
            continue
        seen_ids.add(candidate_id)
        candidates.append(
            PdfFormulaCandidate(
                candidate_id=candidate_id,
                page=raw_page,
                bbox=bbox,
                text=text or None,
                chunk_id=chunk_id,
            )
        )
    candidates.sort(key=_formula_candidate_sort_key)
    return candidates[:limit]


def _formula_chunk_bbox(chunk: Mapping[str, Any]) -> tuple[float, float, float, float] | None:
    """Return normalized chunk geometry only when its declared unit is safe."""

    raw_unit = chunk.get("bbox_unit")
    unit = str(getattr(raw_unit, "value", raw_unit) or "").strip().casefold()
    if unit != "normalized_ratio":
        return None
    return _coerce_normalized_formula_bbox(chunk.get("bbox"))


def bind_pdf_formula_candidates_to_chunks(
    candidates: Sequence[PdfFormulaCandidate],
    chunks: Sequence[Mapping[str, Any]],
) -> list[PdfFormulaCandidate]:
    """Attach the best same-page chunk id without mutating either input."""

    output: list[PdfFormulaCandidate] = []
    for candidate in candidates:
        best: tuple[float, str] | None = None
        candidate_text = _formula_match_key(candidate.text)
        for chunk in chunks:
            if not isinstance(chunk, Mapping):
                continue
            raw_page = chunk.get("page")
            if (
                isinstance(raw_page, bool)
                or not isinstance(raw_page, int)
                or raw_page != candidate.page
            ):
                continue
            chunk_id = str(chunk.get("chunk_id") or "").strip()[:200]
            if not chunk_id:
                continue
            score = 0.0
            chunk_type = str(chunk.get("chunk_type") or "").strip().casefold()
            if chunk_type in {"formula", "equation"} or _normalize_formula_text(
                chunk.get("equation_latex")
            ):
                score += 5.0
            chunk_bbox = _formula_chunk_bbox(chunk)
            if chunk_bbox is not None:
                overlap = _formula_bbox_overlap_ratio(candidate.bbox, chunk_bbox)
                score += overlap * 20.0
                center_x = candidate.bbox[0] + candidate.bbox[2] / 2.0
                center_y = candidate.bbox[1] + candidate.bbox[3] / 2.0
                if (
                    chunk_bbox[0] <= center_x <= chunk_bbox[0] + chunk_bbox[2]
                    and chunk_bbox[1] <= center_y <= chunk_bbox[1] + chunk_bbox[3]
                ):
                    score += 5.0
            chunk_text = _formula_match_key(_formula_chunk_text(chunk))
            if candidate_text and chunk_text:
                if candidate_text == chunk_text:
                    score += 30.0
                elif candidate_text in chunk_text or chunk_text in candidate_text:
                    score += 20.0
            if score >= 10.0 and (
                best is None or score > best[0] or (score == best[0] and chunk_id < best[1])
            ):
                best = (score, chunk_id)
        output.append(
            PdfFormulaCandidate(
                candidate_id=candidate.candidate_id,
                page=candidate.page,
                bbox=candidate.bbox,
                text=candidate.text,
                chunk_id=best[1] if best is not None else candidate.chunk_id,
            )
        )
    return output


def merge_pdf_formula_candidates(
    primary: Sequence[PdfFormulaCandidate],
    detected: Sequence[PdfFormulaCandidate],
    *,
    limit: int = _PDF_FORMULA_SCAN_LIMIT,
) -> list[PdfFormulaCandidate]:
    """Merge candidates, preferring persisted metadata unless it spans atomic boxes."""

    if (
        not isinstance(limit, int)
        or isinstance(limit, bool)
        or limit < 1
        or limit > _PDF_FORMULA_SCAN_LIMIT
    ):
        raise ValueError(f"limit must be between 1 and {_PDF_FORMULA_SCAN_LIMIT}")
    merged: list[PdfFormulaCandidate] = []
    seen_ids: set[str] = set()
    retained_primary = [
        candidate
        for candidate in primary
        if not _formula_candidate_has_atomic_replacements(candidate, detected)
    ]
    for candidate in [*retained_primary, *detected]:
        if candidate.candidate_id in seen_ids:
            continue
        if any(_formula_candidates_are_duplicates(existing, candidate) for existing in merged):
            continue
        seen_ids.add(candidate.candidate_id)
        merged.append(candidate)
    merged.sort(key=_formula_candidate_sort_key)
    return merged[:limit]


def _is_plausible_visual_block(bbox: list[float], page_rect: Any) -> bool:
    """Filter icons, logos, and separators that should not become evidence."""

    try:
        page_width = float(getattr(page_rect, "width", 0.0) or 0.0)
        page_height = float(getattr(page_rect, "height", 0.0) or 0.0)
    except (TypeError, ValueError):
        return False
    if page_width <= 0.0 or page_height <= 0.0:
        return False
    width_px = bbox[2] * page_width
    height_px = bbox[3] * page_height
    if width_px < _PDF_VISUAL_MIN_WIDTH or height_px < _PDF_VISUAL_MIN_HEIGHT:
        return False
    if bbox[2] >= 0.92 and bbox[3] >= 0.88:
        return False
    return bbox[2] * bbox[3] >= _PDF_VISUAL_MIN_AREA_RATIO


def _is_plausible_caption_crop_bbox(bbox: list[float] | None) -> bool:
    """Return whether a caption-neighbor crop is local enough for evidence."""

    if bbox is None or len(bbox) < 4:
        return False
    _x, _y, width, height = bbox
    if width >= 0.92 and height >= 0.88:
        return False
    if width < 0.18 or height < 0.08:
        return False
    if width * height < 0.018:
        return False
    aspect = width / max(height, 1e-6)
    return 0.12 <= aspect <= 8.0


def _nearest_visual_for_caption(
    caption_bbox: list[float] | None,
    page_images: list[_PdfVisualImage],
) -> _PdfVisualImage | None:
    """Return the nearest adjacent image block above or below a caption."""

    if caption_bbox is None or not page_images:
        return None
    caption_left = caption_bbox[0]
    caption_top = caption_bbox[1]
    caption_right = caption_bbox[0] + caption_bbox[2]
    caption_bottom = caption_bbox[1] + caption_bbox[3]
    caption_center_x = caption_bbox[0] + caption_bbox[2] / 2.0
    max_vertical_gap = 0.075
    min_horizontal_overlap = 0.25
    best: tuple[float, _PdfVisualImage] | None = None
    for image in page_images:
        image_left = image.bbox[0]
        image_top = image.bbox[1]
        image_right = image.bbox[0] + image.bbox[2]
        image_bottom = image.bbox[1] + image.bbox[3]
        overlap = max(0.0, min(caption_right, image_right) - max(caption_left, image_left))
        narrow_width = max(0.0001, min(caption_bbox[2], image.bbox[2]))
        horizontal_overlap = overlap / narrow_width
        if horizontal_overlap < min_horizontal_overlap:
            continue
        if caption_bottom <= image_top:
            vertical_gap = image_top - caption_bottom
        elif image_bottom <= caption_top:
            vertical_gap = caption_top - image_bottom
        else:
            vertical_gap = 0.0
        if vertical_gap > max_vertical_gap:
            continue
        image_center_x = image.bbox[0] + image.bbox[2] / 2.0
        horizontal_distance = abs(caption_center_x - image_center_x)
        score = vertical_gap * 4.0 + horizontal_distance
        if best is None or score < best[0]:
            best = (score, image)
    return best[1] if best is not None else None


def _union_normalized_bbox(a: list[float] | None, b: list[float] | None) -> list[float] | None:
    """Return the bounding union of two normalized [x, y, w, h] bboxes."""

    if a is None:
        return list(b) if b is not None else None
    if b is None:
        return list(a)
    x0 = min(a[0], b[0])
    y0 = min(a[1], b[1])
    x1 = max(a[0] + a[2], b[0] + b[2])
    y1 = max(a[1] + a[3], b[1] + b[3])
    return [round(x0, 6), round(y0, 6), round(max(0.0, x1 - x0), 6), round(max(0.0, y1 - y0), 6)]


def _crop_bbox_with_caption(
    visual_bbox: list[float] | None,
    caption_bbox: list[float] | None,
) -> list[float] | None:
    """Union a resolved visual region with its caption so crops carry the label.

    Keeps the visual side anchored and clamps overall height so a caption-bound
    crop stays a local evidence tile rather than a full-page screenshot.
    """

    if visual_bbox is None:
        return None
    merged = _union_normalized_bbox(visual_bbox, caption_bbox)
    if merged is None:
        return None
    max_crop_height = 0.72
    if merged[3] > max_crop_height and caption_bbox is not None:
        visual_top = visual_bbox[1]
        visual_bottom = visual_bbox[1] + visual_bbox[3]
        caption_top = caption_bbox[1]
        caption_bottom = caption_bbox[1] + caption_bbox[3]
        caption_center = caption_bbox[1] + caption_bbox[3] / 2.0
        # The caption must always survive the clamp: readers judge relevance
        # from it. Trim the figure side, never the caption band.
        if caption_center >= visual_bottom:
            # Caption below the figure: anchor the caption bottom, trim the top.
            new_top = max(0.0, caption_bottom - max_crop_height)
            new_top = min(new_top, visual_top)
            merged = [merged[0], new_top, merged[2], caption_bottom - new_top]
        else:
            # Caption above the figure: anchor the caption top, trim the bottom.
            new_top = min(caption_top, visual_top)
            new_bottom = min(visual_bottom, new_top + max_crop_height)
            new_bottom = max(new_bottom, caption_bottom)
            merged = [merged[0], new_top, merged[2], new_bottom - new_top]
    return merged


def _drawing_union_bbox_for_caption(
    page: Any,
    caption_bbox: list[float] | None,
    *,
    kind: str,
) -> list[float] | None:
    """Frame a vector figure via the union of drawings adjacent to a caption.

    Handles figures drawn as vector paths (no embedded raster image and no
    detected table grid), where text-gap heuristics would otherwise crop blank
    whitespace between the caption and body text.
    """

    if caption_bbox is None or kind != "figure":
        return None
    try:
        drawings = page.get_drawings()
        page_width = float(getattr(page.rect, "width", 0.0) or 0.0)
        page_height = float(getattr(page.rect, "height", 0.0) or 0.0)
    except (AttributeError, RuntimeError, TypeError, ValueError):
        return None
    if not isinstance(drawings, list) or page_width <= 0.0 or page_height <= 0.0:
        return None

    caption_top = caption_bbox[1] * page_height
    caption_bottom = (caption_bbox[1] + caption_bbox[3]) * page_height
    col_left = (caption_bbox[0] - 0.065) * page_width
    col_right = (caption_bbox[0] + caption_bbox[2] + 0.065) * page_width
    min_side_px = 2.0

    def union_on_side(*, above: bool) -> tuple[list[float] | None, float]:
        x0 = y0 = float("inf")
        x1 = y1 = float("-inf")
        area = 0.0
        for drawing in drawings:
            rect = drawing.get("rect") if isinstance(drawing, dict) else None
            if rect is None:
                continue
            rx0, ry0, rx1, ry1 = float(rect[0]), float(rect[1]), float(rect[2]), float(rect[3])
            if (rx1 - rx0) < min_side_px or (ry1 - ry0) < min_side_px:
                continue
            if rx1 <= col_left or rx0 >= col_right:
                continue
            if above and ry1 > caption_top + 2.0:
                continue
            if not above and ry0 < caption_bottom - 2.0:
                continue
            x0, y0, x1, y1 = min(x0, rx0), min(y0, ry0), max(x1, rx1), max(y1, ry1)
            area += (rx1 - rx0) * (ry1 - ry0)
        if x1 <= x0 or y1 <= y0:
            return None, 0.0
        return _normalized_bbox((x0, y0, x1, y1), page.rect), area

    above_bbox, above_area = union_on_side(above=True)
    below_bbox, below_area = union_on_side(above=False)
    best = above_bbox if above_area >= below_area else below_bbox
    if best is None:
        return None
    if not _is_plausible_caption_crop_bbox(best):
        return None
    return best


def _expanded_caption_x_range(caption_bbox: list[float]) -> tuple[float, float]:
    """Infer a conservative visual-column x range from a caption bbox."""

    left = caption_bbox[0]
    right = caption_bbox[0] + caption_bbox[2]
    center = caption_bbox[0] + caption_bbox[2] / 2.0
    if caption_bbox[2] < 0.35:
        left = center - 0.40
        right = center + 0.40
    else:
        left -= 0.035
        right += 0.035
    left = max(0.02, left)
    right = min(0.98, right)
    if right - left < 0.18:
        left = max(0.02, center - 0.12)
        right = min(0.98, center + 0.12)
    return left, right


def _caption_neighbor_crop_bbox(
    caption_bbox: list[float] | None,
    text_bboxes: list[list[float]],
    *,
    kind: str,
) -> list[float] | None:
    """Infer a local crop adjacent to a caption when no embedded image exists."""

    if caption_bbox is None:
        return None

    left, right = _expanded_caption_x_range(caption_bbox)
    caption_top = caption_bbox[1]
    caption_bottom = caption_bbox[1] + caption_bbox[3]
    max_crop_height = 0.48

    above_bottoms = [
        bbox[1] + bbox[3]
        for bbox in text_bboxes
        if bbox is not caption_bbox and bbox[1] + bbox[3] <= caption_top - 0.012
    ]
    below_tops = [
        bbox[1]
        for bbox in text_bboxes
        if bbox is not caption_bbox and bbox[1] >= caption_bottom + 0.012
    ]

    def above() -> list[float] | None:
        bottom = caption_top - 0.008
        top = (max(above_bottoms) + 0.008) if above_bottoms else max(0.02, bottom - 0.32)
        top = max(0.02, min(top, bottom - 0.08))
        if bottom - top > max_crop_height:
            top = bottom - max_crop_height
        return [round(left, 6), round(top, 6), round(right - left, 6), round(bottom - top, 6)]

    def below() -> list[float] | None:
        top = caption_bottom + 0.008
        bottom = (min(below_tops) - 0.008) if below_tops else min(0.98, top + 0.32)
        bottom = min(0.98, max(bottom, top + 0.08))
        if bottom - top > max_crop_height:
            bottom = top + max_crop_height
        return [round(left, 6), round(top, 6), round(right - left, 6), round(bottom - top, 6)]

    candidates = [below(), above()] if kind == "table" else [above(), below()]
    for candidate in candidates:
        if _is_plausible_caption_crop_bbox(candidate):
            return candidate
    return None


def _horizontal_bbox_overlap_ratio(a: list[float], b: list[float]) -> float:
    """Return horizontal overlap relative to the narrower normalized bbox."""

    overlap = max(0.0, min(a[0] + a[2], b[0] + b[2]) - max(a[0], b[0]))
    return overlap / max(0.0001, min(a[2], b[2]))


def _looks_like_borderless_table_fragment(text: str, bbox: list[float]) -> bool:
    """Return whether a compact text block has a deterministic table-cell shape."""

    if bbox[2] < 0.28 or bbox[3] <= 0.0 or bbox[3] > 0.12:
        return False
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    if len(lines) >= 3:
        return True
    return str(text or "").count("\t") >= 2 or len(re.findall(r"\s{2,}", str(text or ""))) >= 2


def _borderless_table_bbox_from_text_candidates(
    caption_index: int,
    text_candidates: list[tuple[dict[str, Any], str, list[float] | None]],
) -> list[float] | None:
    """Frame a borderless table from compact text blocks after its caption.

    PyMuPDF cannot detect tables that have no ruling lines. This fallback is
    deliberately narrow: it accepts only immediately adjacent, horizontally
    overlapping, multi-cell text blocks and stops at the first prose-like or
    geometrically distant block.
    """

    if caption_index < 0 or caption_index >= len(text_candidates):
        return None
    caption_bbox = text_candidates[caption_index][2]
    if caption_bbox is None:
        return None

    fragments_bbox: list[float] | None = None
    previous_bbox = caption_bbox
    fragment_count = 0
    for _raw_block, text, bbox in text_candidates[caption_index + 1 : caption_index + 5]:
        if bbox is None or not _looks_like_borderless_table_fragment(text, bbox):
            break
        vertical_gap = bbox[1] - (previous_bbox[1] + previous_bbox[3])
        if vertical_gap < -0.008 or vertical_gap > 0.035:
            break
        if _horizontal_bbox_overlap_ratio(previous_bbox, bbox) < 0.55:
            break
        candidate_union = _union_normalized_bbox(fragments_bbox, bbox)
        if candidate_union is None or candidate_union[3] > 0.24:
            break
        fragments_bbox = candidate_union
        previous_bbox = bbox
        fragment_count += 1

    if fragment_count == 0 or fragments_bbox is None:
        return None
    if fragments_bbox[2] < max(0.35, caption_bbox[2] * 0.75):
        return None

    merged = _union_normalized_bbox(caption_bbox, fragments_bbox)
    if merged is None:
        return None
    horizontal_padding = 0.012
    vertical_padding = 0.008
    left = max(0.0, merged[0] - horizontal_padding)
    top = max(0.0, merged[1] - vertical_padding)
    right = min(1.0, merged[0] + merged[2] + horizontal_padding)
    bottom = min(1.0, merged[1] + merged[3] + vertical_padding)
    padded = [
        round(left, 6),
        round(top, 6),
        round(right - left, 6),
        round(bottom - top, 6),
    ]
    if padded[2] >= 0.96 or padded[3] < 0.04 or padded[3] > 0.30:
        return None
    return padded


def _table_regions_from_page(page: Any, page_index: int) -> list[_PdfVisualRegion]:
    """Return detected table regions from PyMuPDF when available."""

    regions: list[_PdfVisualRegion] = []
    try:
        finder = page.find_tables()
    except (AttributeError, RuntimeError, TypeError, ValueError):
        return regions
    tables = getattr(finder, "tables", None)
    if not isinstance(tables, list):
        return regions
    for table in tables:
        bbox = _normalized_bbox(getattr(table, "bbox", None), page.rect)
        if bbox is None or not _is_plausible_caption_crop_bbox(bbox):
            continue
        regions.append(_PdfVisualRegion(page=page_index, bbox=bbox, kind="table"))
    return regions


def _nearest_region_for_caption(
    caption_bbox: list[float] | None,
    regions: list[_PdfVisualRegion],
    *,
    kind: str,
) -> _PdfVisualRegion | None:
    """Return the nearest detected figure/table region for a caption."""

    if caption_bbox is None or not regions:
        return None
    caption_left = caption_bbox[0]
    caption_top = caption_bbox[1]
    caption_right = caption_bbox[0] + caption_bbox[2]
    caption_bottom = caption_bbox[1] + caption_bbox[3]
    caption_center_x = caption_bbox[0] + caption_bbox[2] / 2.0
    max_vertical_gap = 0.18 if kind == "table" else 0.075
    min_horizontal_overlap = 0.20
    best: tuple[float, _PdfVisualRegion] | None = None
    for region in regions:
        if region.kind != kind:
            continue
        region_left = region.bbox[0]
        region_top = region.bbox[1]
        region_right = region.bbox[0] + region.bbox[2]
        region_bottom = region.bbox[1] + region.bbox[3]
        overlap = max(0.0, min(caption_right, region_right) - max(caption_left, region_left))
        narrow_width = max(0.0001, min(caption_bbox[2], region.bbox[2]))
        horizontal_overlap = overlap / narrow_width
        if horizontal_overlap < min_horizontal_overlap:
            continue
        if caption_bottom <= region_top:
            vertical_gap = region_top - caption_bottom
        elif region_bottom <= caption_top:
            vertical_gap = caption_top - region_bottom
        else:
            vertical_gap = 0.0
        if vertical_gap > max_vertical_gap:
            continue
        region_center_x = region.bbox[0] + region.bbox[2] / 2.0
        horizontal_distance = abs(caption_center_x - region_center_x)
        score = vertical_gap * 4.0 + horizontal_distance
        if best is None or score < best[0]:
            best = (score, region)
    return best[1] if best is not None else None


def _rect_from_normalized_bbox(page: Any, bbox: list[float]) -> Any | None:
    """Convert a normalized bbox to a PyMuPDF Rect for rendering."""

    try:
        import pymupdf

        width = float(getattr(page.rect, "width", 0.0) or 0.0)
        height = float(getattr(page.rect, "height", 0.0) or 0.0)
        if width <= 0.0 or height <= 0.0:
            return None
        x, y, w, h = bbox
        rect = new_pymupdf_rect(
            pymupdf,
            x * width,
            y * height,
            (x + w) * width,
            (y + h) * height,
        )
        if pymupdf_rect_area(rect) <= 0:
            return None
        return rect
    except (ImportError, TypeError, ValueError, AttributeError):
        return None


def _rect_overlap_area(a: list[float], b: list[float]) -> float:
    """Return the overlap area of two normalized [x, y, w, h] bboxes."""

    ox = max(0.0, min(a[0] + a[2], b[0] + b[2]) - max(a[0], b[0]))
    oy = max(0.0, min(a[1] + a[3], b[1] + b[3]) - max(a[1], b[1]))
    return ox * oy


def _crop_has_visual_anchor(
    page: Any,
    crop_bbox: list[float],
    raster_bboxes: list[list[float]],
) -> bool:
    """Return whether a geometric fallback crop actually covers real ink.

    Guards the caption-neighbor fallback, the only crop path with no intrinsic
    visual anchor: a caption whose figure lives on another page (cross-page) or
    is followed by plain prose would otherwise render a false evidence tile of
    text pixels. Vector drawings and embedded image blocks overlapping the crop
    are the structural proof that a figure/table body is present on this page.
    """

    crop_area = crop_bbox[2] * crop_bbox[3]
    if crop_area <= 0.0:
        return False

    visual_area = 0.0
    for raster in raster_bboxes:
        visual_area += _rect_overlap_area(crop_bbox, raster)
    if visual_area / crop_area >= 0.05:
        return True

    try:
        drawings = page.get_drawings()
        page_width = float(getattr(page.rect, "width", 0.0) or 0.0)
        page_height = float(getattr(page.rect, "height", 0.0) or 0.0)
    except (AttributeError, RuntimeError, TypeError, ValueError):
        return False
    if not isinstance(drawings, list) or page_width <= 0.0 or page_height <= 0.0:
        return False

    min_side_px = 2.0
    for drawing in drawings:
        rect = drawing.get("rect") if isinstance(drawing, dict) else None
        if rect is None:
            continue
        try:
            rx0, ry0, rx1, ry1 = float(rect[0]), float(rect[1]), float(rect[2]), float(rect[3])
        except (TypeError, ValueError, IndexError):
            continue
        if (rx1 - rx0) < min_side_px or (ry1 - ry0) < min_side_px:
            continue
        stroke_bbox = _normalized_bbox((rx0, ry0, rx1, ry1), page.rect)
        if stroke_bbox is None:
            continue
        visual_area += _rect_overlap_area(crop_bbox, stroke_bbox)
        if visual_area / crop_area >= 0.05:
            return True
    return False


def _render_caption_crop_asset(
    page: Any,
    bbox: list[float],
    output_path: Path,
) -> bool:
    """Render a local caption-neighbor crop as a browser-safe PNG."""

    clip = _rect_from_normalized_bbox(page, bbox)
    if clip is None:
        return False
    try:
        import pymupdf

        pixmap = page.get_pixmap(
            matrix=new_pymupdf_matrix(pymupdf, 2, 2),
            alpha=False,
            clip=clip,
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        pixmap.save(str(output_path))
        return output_path.is_file()
    except (OSError, RuntimeError, TypeError, ValueError, AttributeError):
        return False


def _materialize_pdf_visual_image(image: _PdfVisualImage, project_data_root: Path) -> bool:
    """Write an embedded PDF image only after a caption owns it."""

    output_path = project_data_root / image.path
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        if not output_path.exists():
            output_path.write_bytes(image.data)
        return output_path.is_file()
    except (OSError, RuntimeError, TypeError, ValueError):
        return False


def _caption_ref_if_likely(
    text: str,
    source_digest: str,
    *,
    has_adjacent_visual: bool,
) -> _PdfFigureTableRef | None:
    """Return a caption ref without mistaking body mentions for captions."""

    match = _PDF_FIGURE_TABLE_REF_RE.match(text or "")
    if match is None:
        return None
    suffix = str(text or "")[match.end() :].strip()
    if _PDF_BODY_REF_VERB_RE.match(suffix) or _PDF_BODY_REF_CJK_VERB_RE.match(suffix):
        return None
    if has_adjacent_visual:
        return _ref_from_match(match, source_digest)
    if not suffix:
        return None
    if suffix[0] in ".．:：;；,-–—)）":
        return _ref_from_match(match, source_digest)
    if len(str(text or "").strip()) <= 180:
        return _ref_from_match(match, source_digest)
    return None


def _extract_pymupdf_visual_blocks(
    filename: str,
    source_path: Path,
    *,
    project_id: str | None = None,
    project_data_root: Path | None = None,
) -> list[StructuredBlock] | None:  # type: ignore[valid-type]
    """Extract text blocks and real embedded image assets from a PDF.

    Args:
        filename: Display filename used to build stable project asset paths.
        source_path: Existing PDF path.
        project_id: Project id whose data root owns extracted image files.
        project_data_root: Test hook or pre-resolved project data directory.

    Returns:
        Structured blocks when PyMuPDF can parse the PDF, otherwise ``None``.
        Image paths are project-relative and point to extracted embedded image
        bytes, never full-page preview screenshots.
    """

    if StructuredBlock is None:
        return None
    if not isinstance(source_path, Path) or not source_path.is_file():
        raise ValueError("source_path must be an existing file")
    if project_data_root is None and not str(project_id or "").strip():
        return None

    try:
        import pymupdf
    except ImportError:
        return None

    if project_data_root is None:
        try:
            if TYPE_CHECKING:
                from literature_assistant.core.project_paths import project_data_path
            else:
                from project_paths import project_data_path

            project_data_root = project_data_path(str(project_id))
        except (OSError, RuntimeError, ValueError) as exc:
            _LOGGER.warning("pdf_visual_asset_root_unavailable filename=%s err=%s", filename, exc)
            return None

    source_digest = hashlib.sha1(str(source_path.resolve()).encode("utf-8")).hexdigest()[:12]
    doc_segment = _safe_asset_segment(Path(filename).stem)
    relative_dir = Path("figure_assets") / "extracted" / f"{doc_segment}-{source_digest}"
    blocks: list[StructuredBlock] = []

    try:
        with open_pymupdf_document(pymupdf, str(source_path)) as doc:
            for page_index, page in enumerate(doc, start=1):
                page_dict = page.get_text("dict", sort=True)
                raw_blocks = page_dict.get("blocks") if isinstance(page_dict, dict) else None
                if not isinstance(raw_blocks, list):
                    continue

                page_images: list[_PdfVisualImage] = []
                page_raster_bboxes: list[list[float]] = []
                text_candidates: list[tuple[dict[str, Any], str, list[float] | None]] = []
                image_index = 0

                for raw_block in raw_blocks:
                    if not isinstance(raw_block, dict):
                        continue
                    block_type = raw_block.get("type")
                    bbox = _normalized_bbox(raw_block.get("bbox"), page.rect)
                    if block_type == 1 and bbox is not None:
                        # Track every non-trivial raster region (including a
                        # full-page scan) as ink evidence for the fallback gate,
                        # even when it is not a materializable primary asset.
                        if bbox[2] >= 0.04 and bbox[3] >= 0.04:
                            page_raster_bboxes.append(bbox)
                    if block_type == 1 and bbox is not None and _is_plausible_visual_block(bbox, page.rect):
                        image_bytes = raw_block.get("image")
                        if not isinstance(image_bytes, (bytes, bytearray)):
                            continue
                        asset = _browser_image_asset(image_bytes, raw_block.get("ext"))
                        if asset is None:
                            continue
                        image_index += 1
                        relative_path = (
                            relative_dir / f"p{page_index:04d}_img{image_index:03d}.{asset.extension}"
                        ).as_posix()
                        page_images.append(
                            _PdfVisualImage(
                                page=page_index,
                                bbox=bbox,
                                path=relative_path,
                                data=asset.data,
                                center_y=bbox[1] + bbox[3] / 2.0,
                            )
                        )
                        continue

                    if block_type == 0:
                        text = _text_from_pymupdf_block(raw_block)
                        if text:
                            text_candidates.append((raw_block, text, bbox))

                text_bboxes = [bbox for _raw_block, _text, bbox in text_candidates if bbox is not None]
                page_regions = _table_regions_from_page(page, page_index)
                caption_paths: set[str] = set()
                caption_crop_index = 0

                for text_index, (_raw_block, text, bbox) in enumerate(text_candidates, start=1):
                    text_refs = _figure_table_refs(text, source_digest)
                    first_ref = text_refs[0] if text_refs else None
                    table_visual_bbox: list[float] | None = None
                    if first_ref is not None and first_ref.kind == "table":
                        table_region = _nearest_region_for_caption(
                            bbox,
                            page_regions,
                            kind="table",
                        )
                        table_visual_bbox = table_region.bbox if table_region is not None else None
                        if table_visual_bbox is None:
                            table_visual_bbox = _borderless_table_bbox_from_text_candidates(
                                text_index - 1,
                                text_candidates,
                            )
                    nearest = (
                        _nearest_visual_for_caption(bbox, page_images)
                        if first_ref is not None and table_visual_bbox is None
                        else None
                    )
                    caption_ref = (
                        _caption_ref_if_likely(
                            text,
                            source_digest,
                            has_adjacent_visual=nearest is not None or table_visual_bbox is not None,
                        )
                        if first_ref is not None
                        else None
                    )
                    image_paths: list[str] = []
                    figure_id: str | None = None
                    table_id: str | None = None
                    block_bbox = bbox
                    block_type = "Text"
                    if caption_ref is not None:
                        block_type = _caption_block_type(caption_ref.kind)
                        if caption_ref.kind == "figure":
                            figure_id = caption_ref.ref_id
                        else:
                            table_id = caption_ref.ref_id
                        if nearest is not None and _materialize_pdf_visual_image(nearest, project_data_root):
                            block_bbox = nearest.bbox
                            image_paths = [nearest.path]
                            caption_paths.add(nearest.path)
                        elif caption_ref.kind in {"figure", "table"}:
                            visual_bbox = table_visual_bbox
                            anchored = table_visual_bbox is not None
                            if visual_bbox is None:
                                region = _nearest_region_for_caption(
                                    bbox,
                                    page_regions,
                                    kind=caption_ref.kind,
                                )
                                visual_bbox = region.bbox if region is not None else None
                                anchored = region is not None
                            if visual_bbox is None and caption_ref.kind == "figure":
                                drawing_bbox = _drawing_union_bbox_for_caption(
                                    page,
                                    bbox,
                                    kind="figure",
                                )
                                if drawing_bbox is not None:
                                    visual_bbox = drawing_bbox
                                    anchored = True
                            if visual_bbox is None:
                                visual_bbox = _caption_neighbor_crop_bbox(
                                    bbox,
                                    text_bboxes,
                                    kind=caption_ref.kind,
                                )
                            # Include the caption text in the crop so readers can
                            # judge relevance directly from the evidence tile.
                            crop_bbox = _crop_bbox_with_caption(visual_bbox, bbox)
                            # Gate the geometry-only fallback: a crop with no real
                            # ink is a cross-page/prose false positive -> keep the
                            # ref as a link instead of rendering text as a figure.
                            if (
                                crop_bbox is not None
                                and not anchored
                                and not _crop_has_visual_anchor(page, crop_bbox, page_raster_bboxes)
                            ):
                                crop_bbox = None
                            if crop_bbox is not None:
                                caption_crop_index += 1
                                relative_path = (
                                    relative_dir / f"p{page_index:04d}_cap{caption_crop_index:03d}.png"
                                ).as_posix()
                                output_path = project_data_root / relative_path
                                if _render_caption_crop_asset(page, crop_bbox, output_path) or output_path.exists():
                                    block_bbox = crop_bbox
                                    image_paths = [relative_path]

                    blocks.append(
                        StructuredBlock(
                            block_id=f"p{page_index}_t{text_index}",
                            page=page_index,
                            bbox=block_bbox,
                            block_type=block_type,
                            markdown=text,
                            bbox_unit=("normalized_ratio" if block_bbox is not None else None),
                            image_paths=image_paths,
                            figure_id=figure_id,
                            table_id=table_id,
                            linked_figure_ids=_dedupe_refs(text_refs, kind="figure")
                            if caption_ref is None
                            else [],
                            linked_table_ids=_dedupe_refs(text_refs, kind="table")
                            if caption_ref is None
                            else [],
                        )
                    )
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        _LOGGER.warning("pdf_visual_blocks_failed filename=%s err=%s", filename, exc)
        return None

    return blocks or None


def _extract_document_content(filename: str, raw: bytes) -> str:
    """Extract textual content from an uploaded document based on file type."""
    content = ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "txt" or ext == "md":
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                content = raw.decode(enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue
    elif ext == "bib":
        for enc in ("utf-8", "latin-1"):
            try:
                content = raw.decode(enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue
    elif ext == "ipynb":
        try:
            notebook = json.loads(raw.decode("utf-8"))
            cells = notebook.get("cells", []) if isinstance(notebook, dict) else []
            parts: list[str] = []

            for idx, cell in enumerate(cells, start=1):
                if not isinstance(cell, dict):
                    continue
                cell_type = str(cell.get("cell_type") or "").strip().lower()
                source = cell.get("source")
                if isinstance(source, list):
                    source_text = "".join(str(x) for x in source)
                else:
                    source_text = str(source or "")
                source_text = source_text.strip()
                if not source_text:
                    continue

                if cell_type == "markdown":
                    parts.append(f"[Notebook Markdown Cell {idx}]\n{source_text}")
                elif cell_type == "code":
                    code_lines = [ln for ln in source_text.splitlines() if ln.strip()][:80]
                    code_excerpt = "\n".join(code_lines)
                    if code_excerpt:
                        parts.append(f"[Notebook Code Cell {idx}]\n{code_excerpt}")

                    outputs = cell.get("outputs", [])
                    if isinstance(outputs, list):
                        output_snippets: list[str] = []
                        for output in outputs:
                            if not isinstance(output, dict):
                                continue
                            # stream output
                            if output.get("output_type") == "stream":
                                text = output.get("text")
                                if isinstance(text, list):
                                    text = "".join(str(x) for x in text)
                                text = str(text or "").strip()
                                if text:
                                    output_snippets.append(text)

                            # execute_result / display_data plain text
                            data = output.get("data")
                            if isinstance(data, dict):
                                plain = data.get("text/plain")
                                if isinstance(plain, list):
                                    plain = "".join(str(x) for x in plain)
                                plain = str(plain or "").strip()
                                if plain:
                                    output_snippets.append(plain)

                        if output_snippets:
                            merged_outputs = "\n".join(output_snippets[:20])
                            parts.append(f"[Notebook Output Cell {idx}]\n{merged_outputs}")

            content = "\n\n".join(parts)
            if not content.strip():
                content = f"[Notebook 文件: {filename}，未提取到可索引内容]"
        except (UnicodeDecodeError, json.JSONDecodeError, TypeError, ValueError) as exc:
            content = f"[Notebook 解析失败: {exc}]"
    elif ext == "pdf":
        try:
            import io
            try:
                import pymupdf  # PyMuPDF (fitz)
                with open_pymupdf_document(
                    pymupdf,
                    stream=raw,
                    filetype="pdf",
                ) as pdf_document:
                    pages = [page.get_text() for page in pdf_document.pages()]
                    content = "\n\n".join(pages)
            except ImportError:
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(raw))
                    pages = [page.extract_text() or "" for page in reader.pages]
                    content = "\n\n".join(pages)
                except ImportError:
                    content = f"[PDF 文件: {filename}，需安装 pymupdf 或 PyPDF2 才能提取文本]"
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            content = f"[PDF 解析失败: {exc}]"
    elif ext in ("docx",):
        try:
            import io
            from docx import Document as DocxDocument
            docx_document = DocxDocument(io.BytesIO(raw))
            content = "\n".join(
                paragraph.text
                for paragraph in docx_document.paragraphs
                if paragraph.text.strip()
            )
        except ImportError:
            content = f"[DOCX 文件: {filename}，需安装 python-docx 才能提取文本]"
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            content = f"[DOCX 解析失败: {exc}]"
    else:
        try:
            content = raw.decode("utf-8")
        except UnicodeDecodeError:
            content = f"[未知格式文件: {filename}]"

    return content


def _extract_document_payload_from_path(
    filename: str,
    source_path: Path,
    *,
    project_id: str | None = None,
    project_data_root: Path | None = None,
) -> ExtractedDocumentPayload:
    """Extract content + optional structured blocks + optional markdown_full.

    Replaces the legacy content-only return with a structured payload. For
    PDFs, the core backend is PyMuPDF (see ``pdf_backends.get_pdf_backend``):

      - ``PyMuPDFBackend`` — byte-level identical to legacy behavior;
        ``blocks`` and ``markdown_full`` are always None.

    Non-PDF formats (DOCX, plaintext, etc.) go through the legacy text-only
    paths; ``blocks`` / ``markdown_full`` are None for those.

    Args:
        filename: Display filename used to choose parser behavior.
        source_path: Existing local file path containing the uploaded bytes.
        project_id: Optional project owning extracted PDF visual assets.
        project_data_root: Optional validated data root used by focused tests.

    Returns:
        ``ExtractedDocumentPayload`` — never raises for the PDF/DOCX branches
        (placeholders are returned as content instead). PDF payloads also carry
        actual parser provenance when the backend package is available.

    Raises:
        TypeError / ValueError: If ``source_path`` is not a Path / not a file.
    """

    if not isinstance(source_path, Path):
        raise TypeError("source_path must be a pathlib.Path")
    if not source_path.is_file():
        raise ValueError(f"source_path is not a file: {source_path}")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # PDF: route through backend abstraction
    if (
        ext == "pdf"
        and get_pdf_backend is not None
        and parse_pdf_with_provenance is not None
    ):
        backend = get_pdf_backend()
        try:
            parse_result = parse_pdf_with_provenance(backend, source_path)
            text, blocks, markdown_full = parse_result.legacy_tuple()
            if blocks is None and isinstance(backend, PyMuPDFBackend):
                blocks = _extract_pymupdf_visual_blocks(
                    filename,
                    source_path,
                    project_id=project_id,
                    project_data_root=project_data_root,
                )
            payload = ExtractedDocumentPayload(
                content=text,
                blocks=blocks,
                markdown_full=markdown_full,
                parser_provenance=parse_result.provenance,
                parser_output_sha256=_document_text_sha256(text),
            )
            if apply_pdf_ocr_if_needed is None:
                return payload
            return cast(ExtractedDocumentPayload, apply_pdf_ocr_if_needed(filename, source_path, payload))
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            _LOGGER.warning(
                "PDF backend %r failed parsing %s: %s; "
                "falling back to PyMuPDF",
                getattr(backend, "name", "?"),
                filename,
                exc,
            )
            if PyMuPDFBackend is not None and not isinstance(
                backend, PyMuPDFBackend  # avoid re-entering same failing backend
            ):
                fallback_result = parse_pdf_with_provenance(
                    PyMuPDFBackend(),
                    source_path,
                )
                return ExtractedDocumentPayload(
                    content=fallback_result.text,
                    blocks=fallback_result.blocks,
                    markdown_full=fallback_result.markdown_full,
                    parser_provenance=fallback_result.provenance.with_prior_attempt(
                        getattr(backend, "name", "unknown")
                    ),
                    parser_output_sha256=_document_text_sha256(fallback_result.text),
                )
            return ExtractedDocumentPayload(content=f"[PDF 解析失败: {exc}]")

    # DOCX: legacy path, no structured output
    if ext == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(source_path))
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            return ExtractedDocumentPayload(content=text)
        except ImportError:
            return ExtractedDocumentPayload(
                content=f"[DOCX 文件: {filename}，需安装 python-docx 才能提取文本]"
            )
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            return ExtractedDocumentPayload(content=f"[DOCX 解析失败: {exc}]")

    # Other formats: delegate to byte-based helper
    return ExtractedDocumentPayload(
        content=_extract_document_content(filename, source_path.read_bytes())
    )


def _extract_document_content_from_path(filename: str, source_path: Path) -> str:
    """Extract textual content from a bounded local source file.

    LEGACY SIGNATURE — kept verbatim for all existing callers. New code
    should use ``_extract_document_payload_from_path`` to access the
    structured blocks and markdown_full produced by optional external parsers.

    Args:
        filename: Display filename used to choose parser behavior.
        source_path: Existing local file path containing the uploaded bytes.

    Returns:
        Extracted text or the same user-facing parser placeholder strings used
        by the byte-based compatibility helper.

    Raises:
        ValueError: If ``source_path`` is not an existing file.
    """

    return _extract_document_payload_from_path(filename, source_path).content


def _normalize_coverage_text(value: str) -> str:
    """Normalize text only for block/full-text coverage comparison."""

    normalized = unicodedata.normalize("NFKC", str(value or ""))
    return re.sub(r"\s+", " ", normalized).strip()


def _is_ocr_content_marker(value: str) -> bool:
    normalized = str(value or "").strip()
    return (
        normalized.casefold() == "[ocr extracted text]"
        or bool(_OCR_CONTENT_MARKER_RE.fullmatch(normalized))
        or normalized.casefold().startswith("[ocr not executed for ")
    )


def _reconcile_document_blocks(
    content: str,
    blocks: Sequence[StructuredBlock] | None,  # type: ignore[valid-type]
) -> list[StructuredBlock] | None:  # type: ignore[valid-type]
    """Append explicit fallback blocks for full-text lines absent from blocks.

    A structured parser can return a useful partial layout tree while its
    normalized full text still contains additional pages or paragraphs. The
    structured chunk path must keep those uncovered lines searchable. Locator
    fields deliberately degrade to page ``0`` and ``bbox=None`` because
    inventing page geometry would produce misleading citations.

    Args:
        content: Authoritative extracted full text.
        blocks: Optional structured parser blocks.

    Returns:
        The original no-block sentinel, or a copied block list followed by
        ``FullTextFallback`` blocks for uncovered text groups.
    """

    if not isinstance(content, str):
        raise TypeError("content must be a string")
    if not blocks:
        return None if blocks is None else []
    if StructuredBlock is None:
        return list(blocks)

    reconciled = list(blocks)
    exact_line_capacity: Counter[str] = Counter()
    normalized_block_parts: list[str] = []
    for block in reconciled:
        markdown = str(getattr(block, "markdown", "") or "")
        normalized_markdown = _normalize_coverage_text(markdown)
        if normalized_markdown:
            normalized_block_parts.append(normalized_markdown)
        for line in markdown.splitlines():
            normalized_line = _normalize_coverage_text(line)
            if normalized_line:
                exact_line_capacity[normalized_line] += 1

    normalized_block_corpus = " ".join(normalized_block_parts)
    consumed_capacity: Counter[str] = Counter()
    substring_capacity: dict[str, int] = {}
    missing_groups: list[list[str]] = []
    current_group: list[str] = []

    def _flush_group() -> None:
        if current_group:
            missing_groups.append(list(current_group))
            current_group.clear()

    for raw_line in content.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            _flush_group()
            continue
        if _is_ocr_content_marker(stripped):
            _flush_group()
            continue

        normalized_line = _normalize_coverage_text(raw_line)
        if not normalized_line:
            _flush_group()
            continue

        available = exact_line_capacity[normalized_line]
        if (
            consumed_capacity[normalized_line] >= available
            and len(normalized_line) >= 24
            and normalized_block_corpus
        ):
            available = max(
                available,
                substring_capacity.setdefault(
                    normalized_line,
                    normalized_block_corpus.count(normalized_line),
                ),
            )
        if consumed_capacity[normalized_line] < available:
            consumed_capacity[normalized_line] += 1
            _flush_group()
            continue
        current_group.append(raw_line.rstrip())
    _flush_group()

    for fallback_index, lines in enumerate(missing_groups, start=1):
        fallback_text = "\n".join(lines).strip()
        if not fallback_text:
            continue
        digest = hashlib.sha256(fallback_text.encode("utf-8")).hexdigest()[:16]
        reconciled.append(
            StructuredBlock(
                block_id=f"fulltext_fallback_{fallback_index}_{digest}",
                page=0,
                bbox=None,
                block_type="FullTextFallback",
                markdown=fallback_text,
            )
        )
    return reconciled


def _truncate_document_content(content: str) -> str:
    """Preserve authoritative extracted text for persistence callers.

    The historical name remains as a compatibility seam for upload and scan
    services. Response previews should be bounded at their response boundary;
    truncating here permanently removed the document tail from truth storage
    and every downstream retrieval index.
    """

    if not isinstance(content, str):
        raise TypeError("content must be a string")
    return content
